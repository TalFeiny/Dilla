"""
Document process service: backend-agnostic single-doc extraction.
Runs full flow in-process: download from storage, extract text (PDF/DOCX),
extract structured data via model_router + JSON schema, update metadata repo.
No subprocess.
"""

import asyncio
import json
import logging
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, Optional

from app.abstractions.document_metadata import DocumentMetadataRepo
from app.abstractions.storage import DocumentBlobStorage

logger = logging.getLogger(__name__)

_UNLINKED = "00000000-0000-0000-0000-000000000000"

# Startup check: pypdf required for PDF extraction
def _check_pypdf() -> None:
    try:
        from pypdf import PdfReader  # noqa: F401
    except ImportError:
        logger.warning(
            "pypdf not installed; PDF text extraction will fail. "
            "Run: pip install pypdf>=4.0.0"
        )


_check_pypdf()  # Run on module load

# OCR availability flag
_OCR_AVAILABLE = False
try:
    import pytesseract  # noqa: F401
    from pdf2image import convert_from_path  # noqa: F401
    _OCR_AVAILABLE = True
except ImportError:
    logger.info(
        "pytesseract/pdf2image not installed; OCR fallback disabled. "
        "Run: pip install pytesseract pdf2image  (and install Tesseract + Poppler system deps)"
    )

# ---------------------------------------------------------------------------
# Legal document extraction schemas — clause-level with parent-child hierarchy
# ---------------------------------------------------------------------------

# Document types that route to legal extraction
LEGAL_DOC_TYPES = {
    "contract", "vendor_contract", "services_agreement", "ip_agreement",
    "term_sheet", "side_letter", "sha", "lpa", "board_consent",
    "employment_agreement", "employment_contract", "nda", "license_agreement", "lease",
    "amendment", "sow", "msa",
    # Commercial contracts (bridge to P&L via contract_pnl_bridge)
    "client_agreement", "client_contract",
    "service_agreement", "product_agreement",
    "software_license",
    # Intercompany (bridge to TP engine via contract_pnl_bridge)
    "intercompany_agreement", "management_agreement",
}

# Base schema — every legal doc gets this
LEGAL_BASE_SCHEMA = {
    "document_type": "string (contract | vendor_contract | services_agreement | ip_agreement | term_sheet | side_letter | sha | lpa | board_consent | employment_agreement | nda | license_agreement | lease | amendment | sow | msa)",
    "parties": "array of objects: [{name: string, role: string (e.g. 'vendor', 'client', 'licensor', 'investor', 'company')}]",
    "effective_date": "string ISO date or null",
    "expiration_date": "string ISO date or null",
    "governing_law": "string or null (jurisdiction)",
    "parent_document_id": "string or null (ID of parent doc this modifies/implements)",
    "supersedes": "string or null (ID of doc this replaces)",
    "modifies_clauses": "array of strings (clause IDs in parent doc that this modifies)",
    "document_lineage": "array of strings (ancestor document IDs, root first)",
    "summary": "string (2-3 sentence overview)",
    "clauses": "array of objects — EXTRACT EVERY MATERIAL CLAUSE as: [{"
        "id: string (hierarchical e.g. '4', '4.1', '4.1.a'), "
        "parent_id: string or null (parent clause id, null for top-level), "
        "children: array of strings (child clause ids), "
        "title: string (clause heading), "
        "text: string (verbatim clause text, first 500 chars), "
        "clause_type: string (e.g. 'liquidation_preference', 'anti_dilution', 'termination', "
            "'auto_renewal', 'indemnification', 'ip_assignment', 'non_compete', 'confidentiality', "
            "'payment_terms', 'liability_cap', 'minimum_commitment', 'data_processing', "
            "'change_of_control', 'most_favored_nation', 'pro_rata', 'board_seat', "
            "'drag_along', 'tag_along', 'exclusivity', 'warranty', 'sla', 'force_majeure'), "
        "flags: array of strings — short natural-language flags describing what is notable, risky, or unusual about THIS clause. "
            "Be specific to the actual content (e.g. '90-day auto-renew with 10-day notice window', 'uncapped indemnity for IP claims', "
            "'no termination for convenience', 'liability cap 5x annual fees — above market', 'exclusivity blocks competing vendors for 3 years'). "
            "Do NOT use generic labels. Every flag must tell the reader something concrete about THIS clause. Empty array if nothing notable, "
        "obligations: array of objects [{party: string, description: string, deadline: string or null, recurring: boolean}], "
        "cross_references: array of objects [{"
            "to_service: string (cap_table | liquidation_waterfall | anti_dilution | pnl | cash_flow), "
            "to_entity: string or null (e.g. 'series_a_preferred'), "
            "field: string (e.g. 'liquidation_pref_multiple'), "
            "value: any, "
            "relationship: string (defines | modifies | constrains | overrides)"
        "}]"
    "}]",
    "erp_attribution": {
        "category": "string or null (revenue | cogs | opex_rd | opex_sm | opex_ga — from ERP taxonomy)",
        "subcategory": "string or null (e.g. 'hosting', 'tools_licenses', 'finance_legal', 'contractor', 'insurance')",
        "hierarchy_path": "string or null (e.g. 'cogs/hosting', 'opex_ga/finance_legal')",
        "annual_value": "number or null (total annual contract value, USD)",
        "payment_frequency": "string or null (monthly | quarterly | annual | one_time)",
        "monthly_amount": "number or null (derived or explicit monthly cost)",
        "committed_periods": "array of strings or null (ISO month periods this hits, e.g. ['2026-01', '2026-02'])",
        "variable_component": {
            "estimated_monthly": "number or null",
            "unit": "string or null (e.g. 'API calls', 'seats', 'hours')",
            "rate": "number or null (per-unit rate)",
        },
    },
    "red_flags": "array of strings — concrete issues, not generic labels. Say what's wrong, whether it's market, and what the impact is. "
        "e.g. 'No liability cap — unlimited exposure on a $50K/yr vendor contract', 'Auto-renew with only 10-day notice — easy to miss exit window', "
        "'Personal guarantee from CEO for equipment lease — unusual for this deal size'",
    "key_dates": "array of objects [{event: string, date: string ISO, auto_action: string or null (e.g. 'auto_renew', 'terminate')}]",
    "value_explanations": "object: { [clause_id]: string } — for EACH clause, write a 1-2 sentence explanation of what the clause actually means "
        "and why it matters. Do NOT just copy the clause text — interpret it. Say whether it's market standard or not, "
        "what the practical impact is, and what the reader should care about. "
        "e.g. '4.1': 'Liability capped at 1x annual fees ($120K). Market for this contract size — no action needed.' "
        "e.g. '7.2': 'Auto-renewal every 12 months with 90-day notice requirement. Aggressive — typical notice is 30 days. Must calendar the exit window.' "
        "e.g. '3.1': '2x non-participating liquidation preference. Series A gets 2x back before common sees anything. Reduces common payout by ~$4M at a $20M exit.'",
}

# Term sheet extension — investment deal terms
TERM_SHEET_SCHEMA = {
    **LEGAL_BASE_SCHEMA,
    "valuation_pre_money": "number or null (USD)",
    "valuation_post_money": "number or null (USD)",
    "investment_amount": "number or null (USD)",
    "round": "string or null (e.g. 'Series A', 'Seed')",
    "liquidation_preference": "string or null (e.g. '1x non-participating', '2x participating')",
    "anti_dilution": "string or null (e.g. 'broad-based weighted average', 'full ratchet')",
    "board_seats": "number or null",
    "board_composition": "string or null",
    "pro_rata_rights": "boolean or null",
    "protective_provisions": "array of strings",
    "drag_along_threshold": "string or null (e.g. '75% of preferred')",
    "option_pool": "number or null (percentage, e.g. 0.10 for 10%)",
    "vesting_schedule": "string or null",
    "dividends": "string or null",
}

# Vendor/services contract extension — ERP-heavy
VENDOR_CONTRACT_SCHEMA = {
    **LEGAL_BASE_SCHEMA,
    "vendor_name": "string or null",
    "service_description": "string (what is being provided)",
    "contract_term_months": "number or null",
    "auto_renewal": "boolean or null",
    "renewal_notice_days": "number or null (days before expiry to give notice)",
    "termination_for_convenience": "boolean or null",
    "termination_notice_days": "number or null",
    "sla_uptime": "string or null (e.g. '99.9%')",
    "sla_penalties": "string or null",
    "data_handling": "string or null (processing, storage, transfer restrictions)",
    "liability_cap": "string or null (e.g. '12 months of fees')",
    "insurance_requirements": "string or null",
}

# IP/Tech agreement extension
IP_AGREEMENT_SCHEMA = {
    **LEGAL_BASE_SCHEMA,
    "ip_type": "string or null (patent | copyright | trade_secret | trademark | software)",
    "ownership": "string (work_for_hire | licensed | assigned | joint)",
    "license_type": "string or null (exclusive | non_exclusive | sublicensable)",
    "license_territory": "string or null",
    "license_term": "string or null",
    "royalty_terms": "string or null",
    "source_code_escrow": "boolean or null",
    "reverse_engineering_allowed": "boolean or null",
    "improvements_ownership": "string or null (licensor | licensee | joint)",
}

# Map doc types to their specific schema
LEGAL_SCHEMA_MAP = {
    "term_sheet": TERM_SHEET_SCHEMA,
    "vendor_contract": VENDOR_CONTRACT_SCHEMA,
    "services_agreement": VENDOR_CONTRACT_SCHEMA,
    "msa": VENDOR_CONTRACT_SCHEMA,
    "sow": VENDOR_CONTRACT_SCHEMA,
    "lease": VENDOR_CONTRACT_SCHEMA,
    "ip_agreement": IP_AGREEMENT_SCHEMA,
    "license_agreement": IP_AGREEMENT_SCHEMA,
    "nda": LEGAL_BASE_SCHEMA,
    "employment_agreement": LEGAL_BASE_SCHEMA,
    "employment_contract": LEGAL_BASE_SCHEMA,
    "client_agreement": VENDOR_CONTRACT_SCHEMA,
    "client_contract": VENDOR_CONTRACT_SCHEMA,
    "service_agreement": VENDOR_CONTRACT_SCHEMA,
    "product_agreement": VENDOR_CONTRACT_SCHEMA,
    "software_license": IP_AGREEMENT_SCHEMA,
    "side_letter": LEGAL_BASE_SCHEMA,
    "sha": TERM_SHEET_SCHEMA,
    "lpa": LEGAL_BASE_SCHEMA,
    "board_consent": LEGAL_BASE_SCHEMA,
    "amendment": LEGAL_BASE_SCHEMA,
    "contract": LEGAL_BASE_SCHEMA,
}


def _get_legal_schema(document_type: str) -> dict:
    """Return the appropriate legal extraction schema for a document type."""
    return LEGAL_SCHEMA_MAP.get(document_type, LEGAL_BASE_SCHEMA)


def _empty_legal_extraction() -> dict:
    """Return an empty legal extraction result."""
    return {
        "document_type": None,
        "parties": [],
        "effective_date": None,
        "expiration_date": None,
        "governing_law": None,
        "parent_document_id": None,
        "supersedes": None,
        "modifies_clauses": [],
        "document_lineage": [],
        "summary": "",
        "clauses": [],
        "erp_attribution": {
            "category": None,
            "subcategory": None,
            "hierarchy_path": None,
            "annual_value": None,
            "payment_frequency": None,
            "monthly_amount": None,
            "committed_periods": None,
            "variable_component": None,
        },
        "red_flags": [],
        "key_dates": [],
        "value_explanations": {},
    }


def _normalize_legal_extraction(d: dict) -> dict:
    """Normalize legal extraction — ensure clause hierarchy is well-formed."""
    out = dict(d)
    # Ensure top-level fields exist
    out.setdefault("document_type", None)
    out.setdefault("parties", [])
    out.setdefault("effective_date", None)
    out.setdefault("expiration_date", None)
    out.setdefault("governing_law", None)
    out.setdefault("parent_document_id", None)
    out.setdefault("supersedes", None)
    out.setdefault("modifies_clauses", [])
    out.setdefault("document_lineage", [])
    out.setdefault("summary", "")
    out.setdefault("red_flags", [])
    out.setdefault("key_dates", [])
    out.setdefault("value_explanations", {})

    # Normalize clauses — ensure parent_id/children consistency
    clauses = out.get("clauses")
    if not isinstance(clauses, list):
        out["clauses"] = []
    else:
        clause_ids = {c.get("id") for c in clauses if isinstance(c, dict) and c.get("id")}
        for clause in clauses:
            if not isinstance(clause, dict):
                continue
            clause.setdefault("id", "")
            clause.setdefault("parent_id", None)
            clause.setdefault("children", [])
            clause.setdefault("title", "")
            clause.setdefault("text", "")
            clause.setdefault("clause_type", "other")
            clause.setdefault("flags", [])
            clause.setdefault("obligations", [])
            clause.setdefault("cross_references", [])
            # Validate parent_id references a real clause
            if clause["parent_id"] and clause["parent_id"] not in clause_ids:
                clause["parent_id"] = None
            # Validate children references
            clause["children"] = [c for c in clause["children"] if c in clause_ids]

    # Normalize erp_attribution
    erp = out.get("erp_attribution")
    if not isinstance(erp, dict):
        out["erp_attribution"] = _empty_legal_extraction()["erp_attribution"]
    else:
        erp.setdefault("category", None)
        erp.setdefault("subcategory", None)
        erp.setdefault("hierarchy_path", None)
        erp.setdefault("annual_value", None)
        erp.setdefault("payment_frequency", None)
        erp.setdefault("monthly_amount", None)
        erp.setdefault("committed_periods", None)
        erp.setdefault("variable_component", None)
        # Ensure numeric fields are numeric
        for nk in ("annual_value", "monthly_amount"):
            erp[nk] = _ensure_numeric(erp.get(nk))
        # Auto-derive hierarchy_path
        if erp.get("category") and erp.get("subcategory") and not erp.get("hierarchy_path"):
            erp["hierarchy_path"] = f"{erp['category']}/{erp['subcategory']}"

    # Ensure red_flags is list of strings
    out["red_flags"] = [x for x in (out.get("red_flags") or []) if isinstance(x, str)]
    # Ensure value_explanations is dict
    if not isinstance(out.get("value_explanations"), dict):
        out["value_explanations"] = {}

    return out


def _legal_extraction_prompt(text: str, document_type: str, schema_desc: str, **kwargs) -> tuple:
    """Build system + user prompts for legal document clause extraction."""
    system_prompt = (
        "You are a legal analyst AI specializing in contract clause extraction. "
        "You extract EVERY material clause from legal documents into a structured hierarchy. "
        "You identify parent-child relationships between clauses (Section 4 → 4.1 → 4.1(a)). "
        "You flag what matters about each clause — whether terms are market standard, aggressive, or favorable, and what the concrete impact is. "
        "You identify cross-references to financial services (cap table, P&L, cash flow). "
        "You map ALL commercial contracts to ERP categories (revenue AND cost). "
        "Return ONLY valid JSON matching the schema. No markdown. No explanation."
    )
    user_prompt = (
        f"Extract all clauses and structured data from this {document_type} document.\n\n"
        f"Target JSON schema:\n{schema_desc}\n\n"
        "IMPORTANT:\n"
        "- Extract EVERY material clause with hierarchical IDs (e.g. '1', '1.1', '1.1.a')\n"
        "- Set parent_id for child clauses (e.g. '1.1' has parent_id '1')\n"
        "- Set children arrays (e.g. '1' has children ['1.1', '1.2'])\n"
        "- flags: For each clause, write short concrete flags about what matters. Say whether terms are market/aggressive/favorable and WHY. "
        "e.g. 'Liability cap at 12x annual fees — aggressive, market is 1-2x', 'No termination for convenience — locked in', "
        "'Anti-dilution is broad-based weighted avg — market standard'. Don't use generic labels like 'unfavorable' — say what's unfavorable and whether it's market.\n"
        "- For ALL commercial contracts, ALWAYS populate erp_attribution with category + subcategory:\n"
        "  * Client/customer/revenue contracts → category: 'revenue', subcategory: e.g. 'saas_recurring', 'consulting', 'licensing', 'product_sales', 'subscriptions'\n"
        "  * Vendor contracts, MSAs, SOWs → category: 'cogs' or 'opex_rd'/'opex_sm'/'opex_ga', subcategory from service type\n"
        "  * Leases (office, facility) → category: 'opex_ga' (office/admin) or 'cogs' (production facility), subcategory: 'rent', 'equipment_lease'\n"
        "  * IP/software licenses → category: 'cogs' (core product dependency), 'opex_rd' (dev tools), or 'opex_ga' (business tools), subcategory from software/IP type\n"
        "  * Employment/compensation contracts → category: 'opex_rd' (engineering/product), 'opex_sm' (sales/marketing), 'opex_ga' (admin/finance/legal/HR), subcategory: 'salaries', 'compensation', 'benefits'\n"
        "  * Service agreements → category based on department served, subcategory from service description\n"
        "  * Physical product/supply agreements → category: 'cogs', subcategory: 'materials', 'supplies', 'inventory'\n"
        "- For investment docs: identify cross-references to cap_table, liquidation_waterfall, anti_dilution\n"
        "- Include obligations with party, description, deadline, recurring\n"
        "- Extract verbatim clause text (first 500 chars)\n"
        "- ALWAYS extract document-level dates and numbers:\n"
        "  * effective_date: when the agreement takes effect\n"
        "  * expiration_date: when it expires or terminates\n"
        "  * For term sheets/investment docs: investment_amount, valuation_pre_money, valuation_post_money\n"
        "  * For commercial contracts: erp_attribution.annual_value, erp_attribution.monthly_amount\n"
        "  * parties: extract ALL parties with name and role\n"
        "- value_explanations: write one for EVERY clause — this is critical. Interpret the clause, don't copy it.\n"
    )
    # Inject ERP context hint if provided (from P&L row context)
    erp_category_hint = kwargs.get("erp_category_hint")
    erp_subcategory_hint = kwargs.get("erp_subcategory_hint")
    if erp_category_hint:
        user_prompt += (
            f"\nHINT: This document was uploaded to the '{erp_category_hint}' section of the P&L."
        )
        if erp_subcategory_hint:
            user_prompt += f" Target subcategory: '{erp_subcategory_hint}'."
        user_prompt += " Use this as guidance for erp_attribution.category and subcategory.\n"

    user_prompt += (
        f"\nDocument text:\n---\n{text[:120000]}\n---\n\n"
        "Return only the JSON object, no markdown or explanation."
    )
    return system_prompt, user_prompt


# Document extraction JSON schema (fields we ask the model to return) – used for other doc types
# Note: memos, updates, board decks, board transcripts use signal/memo schemas below
# Flat schema for pitch_deck, other — market_size only in investment_memo
DOCUMENT_EXTRACTION_SCHEMA = {
    "company_name": "string or null",
    "revenue": "number or null (USD)",
    "arr": "number or null (USD)",
    "stage": "string e.g. Seed, Series A",
    "total_funding": "number or null (USD)",
    "valuation": "number or null (USD)",
    "key_metrics": "array of strings",
    "summary": "string",
    "sector": "string or null",
    "target_market": "string or null",
    "business_model": "string or null",
    "red_flags": "array of strings (concerns, risks, concerning language)",
    "value_explanations": "object: { [metric_key]: string } — '\"source quote\" → why → metric change'. e.g. arr: '\"Q3 exceeded plan\" → accelerating sales → ARR up to $1.2M'",
}

# Company-update signal schema – for monthly_update, board_deck, board_transcript
# SERVICE_ALIGNED: business_model, sector, category for valuation/analysis
# new_hires: prefer structured array e.g. [{"role": "Senior PM", "department": "product"}] or string array
COMPANY_UPDATE_SIGNAL_SCHEMA = {
    "company_name": "string or null (optional)",
    "summary": "string (optional)",
    "business_model": "string or null (e.g. SaaS, AI-first, services, rollup—for valuation multiples)",
    "sector": "string or null (e.g. Fintech, Healthcare)",
    "category": "string or null (e.g. saas, ai_first, fintech—for business model detection)",
    "business_updates": {
        "product_updates": "array of strings",
        "achievements": "array of strings",
        "challenges": "array of strings",
        "risks": "array of strings",
        "key_milestones": "array of strings",
        "asks": "array of strings",
        "latest_update": "string (one-line summary of what changed this period — NOT a generic company description)",
        "defensive_language": "array of strings (hedging, caveats, excuses)",
    },
    "operational_metrics": {
        "new_hires": "array of strings or objects. Prefer objects: [{role, department}] e.g. 'Senior PM, product'",
        "headcount": "number or null",
        "customer_count": "number or null",
        "enterprise_customers": "number or null",
        "smb_customers": "number or null",
    },
    "extracted_entities": {
        "competitors_mentioned": "array of strings",
        "industry_terms": "array of strings",
        "partners_mentioned": "array of strings",
    },
    "red_flags": "array of strings (explicit concerns, risks, concerning language)",
    "implications": "array of strings (reading between the lines: inferred items e.g. 'option pool likely expanded given senior product hire')",
    "period_date": "string (ISO date when document period is indicated)",
    "financial_metrics": {
        "arr": "number or null (USD)",
        "revenue": "number or null (USD)",
        "mrr": "number or null (USD)",
        "burn_rate": "number or null (USD)",
        "runway_months": "number or null",
        "cash_balance": "number or null (USD)",
        "gross_margin": "number or null (0-1 or 0-100)",
        "growth_rate": "number or null (e.g. 0.5 for 50%)",
        "customer_count": "number or null",
        "ebitda": "number or null (USD)",
        "ebitda_margin": "number or null (0-1)",
        "operating_income": "number or null (USD)",
        "net_income": "number or null (USD)",
        "capex": "number or null (USD)",
        "fcf": "number or null (USD, free cash flow)",
        "total_debt": "number or null (USD)",
        "interest_expense": "number or null (USD)",
        "working_capital": "number or null (USD, current assets minus current liabilities)",
        "debt_service": "number or null (USD, principal + interest payments per period)",
        "tax_expense": "number or null (USD)",
    },
    "impact_estimates": {
        "_description": "For EACH qualitative signal, estimate material financial impact. FIRST identify business type from the document, THEN use the right framework. Venture startups → ARR/burn/runway fields. PE/operating/traditional companies → revenue/EBITDA/margin/leverage/FCF/working capital fields. Fill WHICHEVER fields match the business.",
        "estimated_arr_impact": "number or null (USD delta — VENTURE/SAAS: e.g. +500000 if 'landed enterprise client')",
        "estimated_burn_impact": "number or null (USD/mo delta — VENTURE: e.g. +60000 if '3 new hires')",
        "estimated_runway_impact": "number or null (months delta — VENTURE: e.g. -3 if 'burn increased')",
        "estimated_revenue_impact": "number or null (USD delta — ALL COMPANIES: e.g. +2000000 if 'won major contract', -1500000 if 'lost key account representing 8% of revenue')",
        "estimated_ebitda_impact": "number or null (USD delta — PE/OPERATING: from EBITDA bridge logic. e.g. +800000 if 'procurement savings captured', -400000 if 'raw material inflation net of pass-through')",
        "estimated_margin_impact": "number or null (ppt delta on EBITDA margin — PE/OPERATING: e.g. +0.02 if 'operating leverage from revenue growth on fixed cost base', -0.03 if 'input cost inflation outpacing pricing')",
        "estimated_fcf_impact": "number or null (USD delta — PE/OPERATING: e.g. +1500000 if 'DSO improved 5 days releasing working capital', -2000000 if 'growth capex for new facility')",
        "estimated_leverage_impact": "number or null (turns delta on Net Debt/EBITDA — PE/OPERATING: e.g. -0.3 if 'mandatory amort + excess cash flow sweep', +0.5 if 'add-on acquisition funded with incremental debt')",
        "estimated_coverage_ratio_impact": "number or null (ratio delta on interest coverage or DSCR — PE/OPERATING: e.g. +0.2 if 'EBITDA growth + lower rate from refi')",
        "estimated_working_capital_impact": "number or null (USD delta — PE/OPERATING: e.g. +500000 if 'DSO -5 days', -800000 if 'inventory build ahead of seasonal demand')",
        "estimated_multiple_impact": "number or null (EV/EBITDA turns delta — PE/OPERATING: e.g. +0.5 if 'recurring revenue mix grew to 60%', -1.0 if 'customer concentration worsened past 30%')",
        "estimated_headcount_impact": "number or null (delta, e.g. +5 if 'hired 5 engineers', -10 if 'RIF'd 10 people')",
        "estimated_cash_impact": "number or null (USD delta, e.g. +5000000 if 'closed Series B' or 'dividend recap' or 'asset disposal proceeds')",
        "estimated_valuation_impact": "number or null (USD delta on enterprise value, e.g. +20000000 if 'raised at 3x last round' or 'EBITDA growth + multiple expansion implies EV +$20M')",
        "estimated_growth_rate_change": "number or null (pct point delta, e.g. +0.1 if 'accelerating', -0.05 if 'slowing')",
        "thesis_tracking": "object or null — PE/OPERATING ONLY: {thesis_status: 'on_track'|'ahead'|'behind'|'at_risk', key_signals: [string], variance_to_plan: string}. Track whether the investment thesis is playing out based on signals in the document.",
        "ebitda_bridge": "object or null — PE/OPERATING ONLY: {prior_period: number|null, volume_price_mix: number|null, cost_savings: number|null, input_cost_changes: number|null, new_initiatives: number|null, one_offs: number|null, fx_impact: number|null, current_period: number|null}. Decompose EBITDA movement into drivers when data allows.",
        "lbo_model_variance": "object or null — PE/OPERATING ONLY: {revenue_vs_model: string|null, ebitda_vs_model: string|null, fcf_vs_model: string|null, leverage_vs_model: string|null, implied_moic_current: number|null, implied_irr_current: number|null}. Track performance vs underwrite assumptions.",
        "covenant_headroom": "object or null — PE/OPERATING ONLY: {leverage_covenant: number|null, leverage_actual: number|null, leverage_headroom: number|null, coverage_covenant: number|null, coverage_actual: number|null, coverage_headroom: number|null, at_risk: boolean}. Flag covenant proximity.",
        "exit_readiness": "object or null — PE/OPERATING ONLY: {readiness_score: string|null, blockers: [string], management_depth: string|null, growth_narrative_strength: string|null, likely_exit_route: string|null}. Assess exit preparedness signals.",
        "impact_reasoning": "object: { [metric_key]: string } — '\"verbatim quote\" → why → metric change'. For venture: customer/burn logic. For PE: EBITDA bridge logic, thesis tracking, LBO model variance, covenant implications. e.g. '\"procurement savings of £600K captured in H1\" → £600K x 1.27 = $762K annualized → EBITDA +$762K, margin +1.5ppt on $50M revenue base'",
    },
    "value_explanations": "object: { [metric_key]: string } — '\"source quote\" → why → metric change'. e.g. venture: arr: '\"hit $1.2M ARR\" → explicit figure → ARR is $1.2M'. PE: ebitda: '\"EBITDA of £4.2M vs £3.8M prior year\" → £4.2M x 1.27 = $5.33M → EBITDA is $5.33M, +11.6% YoY'",
    "pe_operating_metrics": {
        "_description": "PE/OPERATING ONLY. Extract these when the document is from a PE portfolio company or traditional operating business. Leave entire object null for venture startups.",
        "reported_ebitda": "number or null (USD — management-reported EBITDA before any adjustments)",
        "adjusted_ebitda": "number or null (USD — EBITDA after add-backs: one-offs, management fees, run-rate synergies, non-recurring items)",
        "ebitda_addbacks": "object or null — {management_fees: number|null, one_off_costs: number|null, run_rate_synergies: number|null, non_recurring: number|null, other: number|null, total_addbacks: number|null}",
        "covenant_ebitda": "number or null (USD — EBITDA per credit agreement definition, may differ from management EBITDA)",
        "gross_margin_pct": "number or null (0-1)",
        "ebitda_margin_pct": "number or null (0-1)",
        "operating_margin_pct": "number or null (0-1)",
        "net_margin_pct": "number or null (0-1)",
        "revenue_by_segment": "object or null — {segment_name: revenue_usd, ...} — divisional/product line/geographic breakdown",
        "ebitda_by_segment": "object or null — {segment_name: ebitda_usd, ...}",
        "like_for_like_growth": "number or null (organic growth excluding acquisitions, 0-1)",
        "acquisition_revenue_contribution": "number or null (USD — revenue from bolt-ons/add-ons acquired in period)",
        "order_backlog": "number or null (USD — contracted but undelivered revenue, common in manufacturing/services)",
        "book_to_bill": "number or null (ratio — orders received / revenue delivered)",
        "senior_debt": "number or null (USD)",
        "mezzanine_debt": "number or null (USD)",
        "net_debt": "number or null (USD — total debt minus cash)",
        "leverage_ratio": "number or null (Net Debt / EBITDA turns)",
        "interest_coverage": "number or null (EBITDA / interest expense)",
        "fixed_charge_coverage": "number or null ((EBITDA - capex) / (interest + scheduled principal))",
        "dscr": "number or null (debt service coverage ratio)",
        "dso": "number or null (days sales outstanding)",
        "dpo": "number or null (days payable outstanding)",
        "dio": "number or null (days inventory outstanding)",
        "cash_conversion_cycle": "number or null (days — DSO + DIO - DPO)",
        "maintenance_capex": "number or null (USD — sustaining capex)",
        "growth_capex": "number or null (USD — expansion/investment capex)",
        "capex_as_pct_revenue": "number or null (0-1)",
        "roic": "number or null (return on invested capital, 0-1)",
        "roe": "number or null (return on equity, 0-1)",
        "revenue_per_employee": "number or null (USD)",
        "ebitda_per_employee": "number or null (USD)",
        "customer_concentration_top5_pct": "number or null (0-1 — % of revenue from top 5 customers)",
        "recurring_revenue_pct": "number or null (0-1 — % of revenue that is recurring/contracted)",
        "budget_vs_actual": "object or null — {revenue_budget: number|null, revenue_actual: number|null, ebitda_budget: number|null, ebitda_actual: number|null, commentary: string|null}",
        "industry_kpis": "object or null — industry-specific KPIs extracted verbatim. e.g. {utilization_rate: 0.78, bill_rate: 185, same_store_growth: 0.03, bed_occupancy: 0.91, oee: 0.82, yield_rate: 0.95, fill_rate: 0.97}",
    },
    "time_series": "array of objects or null. TWO types of entry can be mixed in the same array:\n"
        "  TYPE 1 — SUMMARY (one per period): {period: 'YYYY-MM', revenue: number|null, cogs: number|null, gross_profit: number|null, opex: number|null, ebitda: number|null, operating_income: number|null, depreciation: number|null, amortization: number|null, interest_expense: number|null, tax_expense: number|null, net_income: number|null, capex: number|null, fcf: number|null, cash_balance: number|null, total_debt: number|null, working_capital: number|null, debt_service: number|null, dividends: number|null, headcount: number|null, arr: number|null, mrr: number|null, burn_rate: number|null, customers: number|null}\n"
        "  TYPE 2 — LINE ITEM (one per line item per period): {period: 'YYYY-MM', subcategory: 'raw_materials', parent_category: 'cogs', amount: 850000}. "
        "Extract EVERY specific line item the business reports as a subcategory entry. These are the raw operating data: direct materials, direct labor, factory overhead, freight, food cost, rent, utilities, clinical supplies — whatever the actual business calls them. "
        "Normalize: lowercase, underscores (e.g. 'Raw Materials' → 'raw_materials'). The more granular the better.\n"
        "  Include BOTH summary entries AND line-item entries. Extract ALL periods.",
}

# Investment memo schema – for investment_memo
INVESTMENT_MEMO_SCHEMA = {
    "company_name": "string or null",
    "investment_date": "string or null (ISO)",
    "round": "string e.g. Series A",
    "valuation_pre_money": "number or null (USD)",
    "deal_terms_summary": "string or null",
    "memo_assumptions": "object (nested key assumptions from the memo)",
    "revenue": "number or null (USD)",
    "arr": "number or null (USD)",
    "runway_months": "number or null",
    "stage": "string or null",
    "total_funding": "number or null (USD)",
    "valuation": "number or null (USD)",
    "key_metrics": "array of strings",
    "summary": "string",
    "sector": "string or null",
    "target_market": "string or null",
    "business_model": "string or null",
    "market_size": {
        "tam_usd": "number or null (USD)",
        "sam_usd": "number or null (USD)",
        "som_usd": "number or null (USD)",
        "tam_description": "string or null",
        "methodology": "string or null",
    },
    "red_flags": "array of strings (concerns, risks)",
    "value_explanations": "object: { [metric_key]: string } — '\"source quote\" → why → metric change'. e.g. arr: '\"Memo states $2M ARR\" → explicit figure → ARR is $2M'",
}


def _iso_now() -> str:
    from datetime import datetime, timezone
    return datetime.now(timezone.utc).isoformat()


def _ocr_pdf(path: str) -> str:
    """
    OCR a PDF using pytesseract + pdf2image.
    Converts each page to an image and runs Tesseract OCR.
    Returns extracted text or empty string on failure.
    """
    if not _OCR_AVAILABLE:
        logger.warning("OCR requested but pytesseract/pdf2image not installed")
        return ""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        logger.info("Running OCR on PDF: %s", path)
        images = convert_from_path(path, dpi=300)
        text_parts: list[str] = []
        for i, img in enumerate(images):
            try:
                page_text = pytesseract.image_to_string(img, lang="eng")
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
            except Exception as e:
                logger.debug("OCR failed on page %d: %s", i + 1, e)
        result = "\n\n".join(text_parts).strip()
        if result:
            logger.info("OCR extracted %d chars from %d pages", len(result), len(images))
        else:
            logger.warning("OCR produced no text from %d pages", len(images))
        return result
    except Exception as e:
        logger.exception("OCR failed for %s: %s", path, e)
        return ""


# Minimum chars to consider pypdf extraction successful (avoids header-only extracts)
_MIN_TEXT_THRESHOLD = 50


def _extract_pdf_tables(path: str) -> str:
    """Extract tables from a PDF using pdfplumber. Returns pipe-delimited text or empty string."""
    try:
        import pdfplumber
    except ImportError:
        logger.debug("pdfplumber not installed; skipping PDF table extraction")
        return ""
    try:
        table_parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables():
                    if not table:
                        continue
                    rows = []
                    for row in table:
                        cells = [str(cell).strip() if cell else "" for cell in row]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        table_parts.append("\n".join(rows))
        return "\n\n".join(table_parts)
    except Exception as e:
        logger.debug("pdfplumber table extraction failed: %s", e)
        return ""


def _text_from_file(path: str, suffix: str) -> str:
    """
    Extract plain text from a file. Supports PDF (pypdf with OCR fallback) and DOCX (python-docx).
    For PDFs: tries pypdf first, falls back to OCR if text is empty or too short.
    Returns empty string for unsupported types or on error.
    """
    path_obj = Path(path)
    if not path_obj.exists():
        logger.warning("_text_from_file: path does not exist %s", path)
        return ""

    ext = (suffix or path_obj.suffix or "").lower().lstrip(".")
    text_parts: list[str] = []

    try:
        if ext in ("pdf",):
            from pypdf import PdfReader
            reader = PdfReader(path)
            for page in reader.pages:
                try:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
                except Exception as e:
                    logger.debug("pypdf page extract: %s", e)
            text = "\n\n".join(text_parts).strip()

            # Extract tables with pdfplumber for structured data
            table_text = _extract_pdf_tables(path)
            if table_text:
                text = text + "\n\n=== TABLES ===\n" + table_text

            # If we got enough text, use it
            if len(text) >= _MIN_TEXT_THRESHOLD:
                return text

            # Otherwise try OCR fallback
            logger.info(
                "pypdf extracted only %d chars (threshold %d), attempting OCR fallback",
                len(text), _MIN_TEXT_THRESHOLD,
            )
            ocr_text = _ocr_pdf(path)
            if ocr_text:
                return ocr_text

            # Return whatever pypdf got (may be empty)
            return text

        if ext in ("docx", "doc"):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(path)
                paragraphs = "\n\n".join(p.text for p in doc.paragraphs if p.text).strip()

                # Extract tables — financial data often lives here
                table_parts: list[str] = []
                for i, table in enumerate(doc.tables):
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        table_parts.append("\n".join(rows))

                if table_parts:
                    return paragraphs + "\n\n=== TABLES ===\n" + "\n\n".join(table_parts)
                return paragraphs
            except ImportError:
                logger.warning("python-docx not installed; cannot extract .docx text")
                return ""

        if ext in ("csv", "xlsx", "xls"):
            return _text_from_spreadsheet(path, ext)

        logger.warning("_text_from_file: unsupported extension %s", ext)
        return ""
    except Exception as e:
        logger.exception("_text_from_file failed for %s: %s", path, e)
        return ""


def _text_from_spreadsheet(path: str, ext: str) -> str:
    """
    Extract text from CSV/XLSX/XLS spreadsheets using pandas.
    Converts each sheet (or the single CSV) into a readable text representation
    with headers, rows, and summary stats — optimized for LLM financial extraction.
    """
    try:
        import pandas as pd
    except ImportError:
        logger.warning("pandas not installed; cannot extract spreadsheet text")
        return ""

    sheets: dict[str, "pd.DataFrame"] = {}

    try:
        if ext == "csv":
            df = pd.read_csv(path, dtype=str, na_filter=False)
            sheets["Sheet1"] = df
        else:
            # xlsx/xls — read all sheets
            try:
                xls = pd.ExcelFile(path, engine="openpyxl" if ext == "xlsx" else "xlrd")
            except ImportError as ie:
                # fallback: try openpyxl for both
                logger.warning("Excel engine import issue (%s), trying openpyxl", ie)
                xls = pd.ExcelFile(path, engine="openpyxl")
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str, na_filter=False)
                if not df.empty:
                    sheets[sheet_name] = df
    except Exception as e:
        logger.exception("Failed to read spreadsheet %s: %s", path, e)
        return ""

    if not sheets:
        return ""

    parts: list[str] = []
    for sheet_name, df in sheets.items():
        # Skip entirely empty sheets
        if df.shape[0] == 0 and df.shape[1] == 0:
            continue

        # Header
        parts.append(f"=== Sheet: {sheet_name} ({df.shape[0]} rows x {df.shape[1]} cols) ===")
        parts.append(f"Columns: {', '.join(str(c) for c in df.columns)}")
        parts.append("")

        # Render rows as tab-separated text (cap at 2000 rows for PE/complex financials)
        max_rows = 2000
        header_line = "\t".join(str(c) for c in df.columns)
        parts.append(header_line)
        for idx, row in df.head(max_rows).iterrows():
            parts.append("\t".join(str(v) for v in row.values))

        if df.shape[0] > max_rows:
            parts.append(f"... ({df.shape[0] - max_rows} more rows omitted)")

        # Numeric summary for columns that look numeric
        numeric_cols = []
        for col in df.columns:
            try:
                numeric_series = pd.to_numeric(df[col].str.replace(r"[,$%£€]", "", regex=True), errors="coerce")
                non_null = numeric_series.dropna()
                if len(non_null) >= 2:
                    numeric_cols.append((col, non_null))
            except Exception:
                pass

        if numeric_cols:
            parts.append("")
            parts.append("Numeric summary:")
            for col_name, series in numeric_cols:
                parts.append(
                    f"  {col_name}: min={series.min():.2f}, max={series.max():.2f}, "
                    f"mean={series.mean():.2f}, sum={series.sum():.2f}"
                )

        parts.append("")

    text = "\n".join(parts).strip()
    logger.info("Extracted %d chars from spreadsheet (%d sheets)", len(text), len(sheets))
    return text


def _get_memo_context_for_company(
    document_repo: DocumentMetadataRepo,
    company_id: str,
    fund_id: Optional[str] = None,
) -> Optional[str]:
    """
    Get a short reference context string from the latest investment memo for this company.
    Used when processing monthly_update/board_deck so extraction can be "relative to memo".
    """
    filters: Dict[str, Any] = {
        "company_id": company_id,
        "document_type": "investment_memo",
        "status": "completed",
    }
    if fund_id:
        filters["fund_id"] = fund_id
    docs = document_repo.list_(filters=filters, limit=20, offset=0)
    if not docs:
        return None
    # Sort by processed_at desc and take latest
    with_ts = [(d, (d.get("processed_at") or "") or "0") for d in docs]
    with_ts.sort(key=lambda x: x[1], reverse=True)
    doc = with_ts[0][0]
    extracted = doc.get("extracted_data") or {}
    if isinstance(extracted, str):
        try:
            extracted = json.loads(extracted)
        except Exception:
            extracted = {}
    if not isinstance(extracted, dict):
        return None
    fm = extracted.get("financial_metrics") or {}
    arr = extracted.get("arr") or fm.get("arr")
    runway = extracted.get("runway_months") or fm.get("runway_months")
    rev = extracted.get("revenue") or fm.get("revenue")
    val = extracted.get("valuation_pre_money") or extracted.get("valuation") or fm.get("valuation")
    processed_at = doc.get("processed_at") or ""
    parts = ["Baseline from investment memo (same company):"]
    if arr is not None:
        parts.append(f" ARR ${arr:,.0f}" if isinstance(arr, (int, float)) else f" ARR {arr}")
    if runway is not None:
        parts.append(f", runway {runway} months")
    if rev is not None:
        parts.append(f", revenue ${rev:,.0f}" if isinstance(rev, (int, float)) else f", revenue {rev}")
    if val is not None:
        parts.append(f", valuation ${val:,.0f}" if isinstance(val, (int, float)) else f", valuation {val}")
    if processed_at:
        parts.append(f", as of {processed_at[:10]}")
    return "".join(parts) if len(parts) > 1 else None


def _signal_first_prompt(text: str, document_type: str, schema_desc: str, memo_context: Optional[str] = None) -> tuple:
    """Build system and user prompt for signal-first extraction (monthly_update / board_deck / board_transcript)."""
    system_prompt = (
        "You are a financial document analyst and transformation engine. Extract structured signals from company updates, board decks, and board transcripts.\n"
        "You handle ALL business types — venture-backed startups, PE-owned portfolio companies, and traditional operating businesses.\n"
        "CRITICAL FIRST STEP: Determine the business type from the document before applying any analytical framework.\n"
        "  - VENTURE/SAAS: Talks about ARR, MRR, burn rate, runway, fundraising rounds, Series A/B/C, CAC/LTV, PMF.\n"
        "  - PE/OPERATING/TRADITIONAL: Talks about EBITDA, margins, leverage, covenants, debt paydown, capex, working capital, "
        "management accounts, board packs with budget vs actual, add-on acquisitions, 100-day plans, value creation plans.\n"
        "  - If unclear, default to the PE/operating framework — it covers more ground.\n\n"
        "You have TWO jobs:\n"
        "1. EXTRACT: Pull explicit numbers into financial_metrics. Venture → ARR, burn, runway. "
        "PE/operating → revenue, EBITDA, margins, capex, FCF, leverage, debt, working capital, coverage ratios. Extract WHICHEVER metrics appear.\n"
        "2. TRANSFORM: Read qualitative prose and estimate numeric impact in impact_estimates. "
        "For PE companies this means: EBITDA bridge analysis, thesis tracking, LBO model variance, covenant headroom, exit readiness — not just 'revenue went up'.\n\n"
        "RULE: impact_estimates MUST have at least 2 non-null numeric values. A document with zero impact estimates is a failure.\n"
        "RULE: impact_reasoning MUST attribute each estimate to a source quote: '\"quote\" → why → metric change'.\n"
        "RULE: For PE companies, ALWAYS attempt to fill thesis_tracking, ebitda_bridge, and covenant_headroom when signals exist.\n"
        "Return a single JSON object. Use null for truly unknown; use empty arrays for missing lists.\n"
        "CURRENCY CONVERSION (always convert to USD before storing any numeric value):\n"
        "- £ (GBP) → multiply by 1.27\n"
        "- € (EUR) → multiply by 1.09\n"
        "- ¥ (JPY) → divide by 154\n"
        "- ₹ (INR) → divide by 84\n"
        "- Note the original currency and amount in value_explanations (e.g. '£2M → $2.54M USD')."
    )
    user_parts = [
        f"Document type: {document_type}.",
        "Extract signals: product_updates, achievements, challenges, risks, asks, defensive_language, key_milestones.",
        "Then operational_metrics: new_hires (prefer objects with role/department e.g. 'Senior PM, product'), headcount, customer_count.",
        "Then extracted_entities: competitors_mentioned, industry_terms, partners_mentioned.",
        "Extract business_model, sector, category when inferable from context (needed for valuation and analysis).",
        "Extract red_flags: array of explicit concerns, risks, or concerning language.",
        "Extract implications: array of 'reading between the lines' items.",
        "",
        "PATH 1 — Explicit numbers → financial_metrics:",
        "Any number stated in the text goes DIRECTLY into financial_metrics.",
        "  Venture examples: 'we hit $1.2M ARR', 'burn is ~$80K/mo', '45 employees'",
        "  PE/operating examples: 'EBITDA of £4.2M', 'leverage at 3.8x', 'DSO improved to 42 days', 'capex £1.1M YTD', 'interest coverage 2.8x'",
        "",
        "PATH 2 — Qualitative signals → impact_estimates:",
        "Most updates are prose with NO explicit numbers. Reason from signal to magnitude using the methodology below.",
        "For each estimate, impact_reasoning MUST follow: '\"verbatim quote\" → why → metric change'.",
        "MANDATORY: At least 2 non-null impact_estimates per document. Rough is better than null.",
        "",
        "========================================================================",
        "=== VENTURE / SAAS IMPACT METHODOLOGY (use when business is a startup) ===",
        "========================================================================",
        "",
        "STEP 1 — CLASSIFY the signal:",
        "  Revenue signals: customer win/loss/expansion, pricing change, new segment, churn, upsell, product launch",
        "  Cost signals: hiring, departures, office moves, vendor changes, raises, layoffs",
        "  Balance sheet signals: fundraise, debt, large purchases, runway changes",
        "  Growth trajectory signals: market expansion, pivot, acceleration/deceleration, PMF indicators",
        "",
        "STEP 2 — ANCHOR to company scale:",
        "  If a BASELINE ANCHOR is provided above, USE IT.",
        "  If no baseline, infer from clues: 10-person Series A ≈ $500K-3M ARR, 200-person post-C ≈ $20-80M ARR.",
        "",
        "STEP 3 — SIZE the impact:",
        "  Revenue: single SMB customer +1-3% ARR | enterprise deal +5-15% ARR | new segment +10-25% yr1 | lost customer -3-10% ARR",
        "  Burn: junior hire +$8-15K/mo | senior hire +$15-30K/mo | departure: reverse",
        "  Growth rate: acceleration +5-15 ppt | deceleration -5-20 ppt",
        "",
        "STEP 4 — CONVERT: proportion x anchored scale = dollar impact. Show math in impact_reasoning.",
        "",
        "VENTURE EXAMPLES:",
        "  'We expanded into enterprise' → new segment at ~$2M ARR baseline, +10-20% yr1 → estimated_arr_impact: +300000",
        "  'Signed LOI with two F500' → 2 x ~$300K ACV, ~60% close rate → estimated_arr_impact: +400000",
        "  'Head of engineering left' → senior eng ~$25K/mo loaded → estimated_burn_impact: -25000",
        "  'Burn is ~$80K/mo' → EXPLICIT: put 80000 in BOTH financial_metrics.burn_rate AND estimated_burn_impact",
        "",
        "=================================================================================",
        "=== PE / OPERATING / TRADITIONAL COMPANY METHODOLOGY (use for non-startups) ===",
        "=================================================================================",
        "",
        "PE board packs and management updates require a fundamentally different analytical framework.",
        "You are analyzing as an investment professional at a PE fund. Think in terms of:",
        "  - VALUE CREATION: What is driving (or destroying) equity value?",
        "  - EBITDA BRIDGE: Decompose EBITDA movements into drivers (volume/price/mix, cost savings, input costs, one-offs, FX)",
        "  - THESIS TRACKING: Is the original investment thesis playing out? Ahead, on track, behind, or at risk?",
        "  - LBO MODEL VARIANCE: How does actual performance compare to the underwrite model?",
        "  - LEVERAGE TRAJECTORY: Is the company de-levering as planned? Covenant headroom?",
        "  - EXIT READINESS: Are signals pointing toward or away from a successful exit?",
        "",
        "STEP 1 — IDENTIFY THE VALUE CREATION LEVERS being discussed:",
        "  Revenue growth: organic (volume, price, mix) vs inorganic (add-ons, bolt-ons)",
        "  Margin expansion: procurement savings, operational efficiency, pricing power, SG&A leverage, mix improvement",
        "  Margin compression: input cost inflation, wage pressure, competitive pricing, regulatory costs, integration costs",
        "  Working capital: DSO/DPO/DIO changes, seasonal patterns, contract term changes, inventory strategy",
        "  Capex: maintenance vs growth capex, facility expansion, technology investment, asset disposals",
        "  Capital structure: debt paydown, refinancing, dividend recaps, add-on financing, covenant amendments",
        "  Multiple drivers: recurring revenue mix, customer diversification, market position, growth profile, ESG",
        "",
        "STEP 2 — ANCHOR to company scale and capital structure:",
        "  If a BASELINE ANCHOR is provided, USE IT for revenue, EBITDA, leverage, debt quantum.",
        "  If no baseline, infer from document signals:",
        "    - 50-employee services business ≈ $5-15M revenue, $1-3M EBITDA",
        "    - 200-employee manufacturer ≈ $30-80M revenue, $5-15M EBITDA",
        "    - 500-employee mid-market company ≈ $80-250M revenue, $15-40M EBITDA",
        "    - References to 'covenant headroom', 'leverage', 'debt service' → leveraged PE deal",
        "    - References to 'budget vs actual', '100-day plan' → recent PE acquisition",
        "  State your assumed scale in impact_reasoning.",
        "",
        "STEP 3 — ANALYZE through the PE lens:",
        "",
        "  A. EBITDA BRIDGE (always attempt when financial data is present):",
        "    Decompose EBITDA change into: volume/price/mix → cost savings → input cost changes → new initiatives → one-offs → FX",
        "    Fill the ebitda_bridge object in impact_estimates.",
        "",
        "  B. REVENUE IMPACT — think in organic growth rate and absolute contribution:",
        "    New contract/customer: estimate annual value from customer type, industry, deal size signals",
        "    Lost customer: estimate revenue at risk, consider replacement timeline",
        "    Price increase: % increase x affected revenue base",
        "    Add-on acquisition: acquired revenue + synergy estimate",
        "    Market/macro effect: industry growth/contraction applied to company's exposure",
        "",
        "  C. MARGIN IMPACT — think in basis points on EBITDA margin:",
        "    Procurement savings: $ saved / revenue base = margin ppt",
        "    Headcount changes: fully loaded cost / revenue base = margin ppt",
        "    Input cost inflation: $ increase / revenue, net of pass-through pricing",
        "    Operating leverage: incremental margin on revenue growth (typically 30-60% for operating cos)",
        "    Integration costs: one-time vs recurring, separate from run-rate margin",
        "",
        "  D. LEVERAGE & COVENANT ANALYSIS:",
        "    Debt paydown from FCF → leverage reduction in turns (debt reduction / EBITDA)",
        "    EBITDA growth → leverage reduction even without paydown",
        "    Add-on debt → leverage increase, check against covenant levels",
        "    Refinancing → interest rate change → coverage ratio change → FCF impact",
        "    ALWAYS check: is covenant headroom tightening or expanding?",
        "",
        "  E. WORKING CAPITAL & FCF:",
        "    DSO change: (days change / 365) x annual revenue = cash released/consumed",
        "    DIO change: (days change / 365) x annual COGS = cash released/consumed",
        "    DPO change: (days change / 365) x annual purchases = cash released/consumed",
        "    Seasonal working capital: don't confuse with structural change",
        "    FCF = EBITDA - cash tax - maintenance capex - cash interest - working capital change",
        "",
        "  F. THESIS TRACKING:",
        "    Map document signals to the likely investment thesis:",
        "      Buy-and-build: Are add-ons being executed? Integration on track? Synergies captured?",
        "      Margin expansion: Are operational improvements materializing? Procurement savings on plan?",
        "      Revenue growth: New markets/products/channels delivering? Organic growth accelerating?",
        "      Management upgrade: New hires performing? Organizational capability improving?",
        "      Digital transformation: Technology investment yielding efficiency? Data-driven decisions?",
        "    Fill thesis_tracking with status and key signals.",
        "",
        "  G. EXIT READINESS SIGNALS:",
        "    Positive: recurring revenue growing, customer diversification improving, strong management bench, clean financials, market tailwinds",
        "    Negative: customer concentration, key-man risk, messy carve-out issues, regulatory overhang, leverage still high",
        "    Fill exit_readiness when signals exist.",
        "",
        "  H. INDUSTRY-SPECIFIC KPIS (extract when present):",
        "    Manufacturing: OEE, yield, scrap rate, book-to-bill, order backlog, capacity utilization",
        "    Services: utilization rate, bill rate, pipeline, win rate, revenue per FTE",
        "    Retail: same-store sales, basket size, footfall, shrinkage, inventory turns",
        "    Healthcare: patient volume, payer mix, collections rate, PMPM, bed occupancy",
        "    Distribution: fill rate, delivery cost/unit, warehouse utilization, fleet efficiency",
        "    Technology (non-SaaS): license revenue, maintenance attach, professional services margin",
        "    Put these in operational_metrics or financial_metrics as appropriate.",
        "",
        "STEP 4 — CONVERT and show your math:",
        "  All impacts must be converted to USD amounts or ppt/turns as appropriate.",
        "  In impact_reasoning: '\"verbatim quote\" → [PE analytical reasoning] → [metric] [direction] $Y or X ppt or X turns'",
        "",
        "PE / OPERATING EXAMPLES (full reasoning chain):",
        "",
        "  'Procurement savings programme delivered £600K in H1'",
        "    Signal: cost savings (margin expansion). Anchor: £600K in H1 = £1.2M annualized = $1.52M.",
        "    estimated_ebitda_impact: +1524000, estimated_margin_impact: +0.025 (assuming ~$60M revenue)",
        "    ebitda_bridge: {cost_savings: 1524000}",
        "    thesis_tracking: {thesis_status: 'on_track', key_signals: ['procurement savings delivering']}",
        "    impact_reasoning: '\"procurement savings £600K in H1\" → annualized £1.2M x 1.27 = $1.52M → EBITDA +$1.52M, +2.5ppt on ~$60M revenue'",
        "",
        "  'Net debt reduced to 3.2x EBITDA from 4.1x at acquisition'",
        "    Signal: de-leveraging (capital structure). Prior 4.1x → current 3.2x = 0.9 turns improvement.",
        "    estimated_leverage_impact: -0.9",
        "    covenant_headroom: {leverage_actual: 3.2, at_risk: false}",
        "    impact_reasoning: '\"net debt 3.2x from 4.1x at acquisition\" → 0.9 turns de-leveraging → equity value accreting, covenant headroom expanding'",
        "",
        "  'Completed bolt-on acquisition of SmithCo for £8M (6x EBITDA)'",
        "    Signal: add-on (inorganic growth). SmithCo EBITDA = £8M/6 = £1.33M = $1.69M. Revenue likely ~$10-12M.",
        "    estimated_revenue_impact: +12000000, estimated_ebitda_impact: +1690000, estimated_leverage_impact: +0.3",
        "    thesis_tracking: {thesis_status: 'on_track', key_signals: ['buy-and-build executing']}",
        "    impact_reasoning: '\"acquired SmithCo for £8M at 6x\" → EBITDA £1.33M = $1.69M, revenue ~$10-12M → platform EBITDA +$1.69M, leverage +~0.3 turns from acquisition debt'",
        "",
        "  'DSO improved from 52 to 45 days'",
        "    Signal: working capital improvement (FCF). Anchor: assume ~$50M revenue.",
        "    Cash release: (7/365) x $50M = $959K. Structural improvement, not seasonal.",
        "    estimated_working_capital_impact: +959000, estimated_fcf_impact: +959000",
        "    impact_reasoning: '\"DSO 52 to 45 days\" → 7 days on ~$50M revenue = $959K cash released → WC +$959K, FCF +$959K'",
        "",
        "  'Key customer (18% of revenue) not renewing contract'",
        "    Signal: customer loss (revenue, concentration). Anchor: 18% of ~$40M = $7.2M at risk.",
        "    estimated_revenue_impact: -7200000, estimated_ebitda_impact: -2500000 (35% incremental margin)",
        "    estimated_multiple_impact: -0.5 (concentration was improving, now worse)",
        "    exit_readiness: {blockers: ['customer concentration worsening'], growth_narrative_strength: 'weakened'}",
        "    impact_reasoning: '\"key customer 18% not renewing\" → $7.2M revenue at risk, ~35% incremental margin → EBITDA -$2.5M, concentration re-risk → multiple -0.5x'",
        "",
        "  'Management presenting 100-day plan update: 3 of 5 initiatives on track, IT integration delayed'",
        "    Signal: thesis tracking (post-acquisition execution). 60% on track, IT delay is a flag.",
        "    thesis_tracking: {thesis_status: 'behind', key_signals: ['3/5 initiatives on track', 'IT integration delayed'], variance_to_plan: '2 initiatives behind schedule'}",
        "    impact_reasoning: '\"3 of 5 on track, IT delayed\" → execution risk on integration synergies, IT delay may push cost savings 2-3 months → thesis tracking: behind'",
        "",
        "  'EBITDA £4.2M vs budget £4.8M, prior year £3.8M'",
        "    Signal: underperformance vs plan but growth YoY. EBITDA = $5.33M, budget miss of £600K = $762K.",
        "    estimated_ebitda_impact: +508000 (vs PY), lbo_model_variance: {ebitda_vs_model: 'behind by $762K (12.5%)', revenue_vs_model: null}",
        "    impact_reasoning: '\"EBITDA £4.2M vs budget £4.8M, PY £3.8M\" → +11% YoY but -12.5% vs plan → growth on track but operational improvement lagging budget'",
        "",
        "For each extracted metric, add to value_explanations: '\"source quote\" → why → metric change'.",
        "For extrapolated values, include the doc excerpt and inference in value_explanations.",
        "",
        "Schema (JSON):",
        schema_desc,
        "",
        "Document text:",
        "---",
        (text[:120000] if len(text) > 120000 else text),
        "---",
        "Return only the JSON object, no markdown or explanation.",
    ]
    if memo_context:
        user_parts.insert(1, f"BASELINE ANCHOR (same company, use for scaling all estimates): {memo_context}")
    user_prompt = "\n".join(user_parts)
    return system_prompt, user_prompt


def _memo_prompt(text: str, schema_desc: str) -> tuple:
    """Build system and user prompt for investment memo extraction."""
    system_prompt = (
        "You are an investment document analyst. Extract structured data from an investment memo — "
        "this could be a venture capital IC memo, a PE deal memo, a credit memo, or an acquisition memo.\n"
        "FIRST identify the deal type from the document:\n"
        "  - VENTURE: talks about round (Seed/Series A/B/C), ARR, burn rate, runway, PMF, CAC/LTV\n"
        "  - PE/BUYOUT: talks about EBITDA, leverage, LBO returns, management buyout, value creation plan, bolt-ons, de-leveraging\n"
        "  - CREDIT/DEBT: talks about coverage ratios, security package, covenants, debt service, collateral\n"
        "  - GROWTH EQUITY: talks about revenue, margins, path to profitability, minority stake, board seat\n"
        "Capture company_name, investment_date, round (or deal_type for PE), valuation_pre_money (or EV for PE), deal_terms_summary, "
        "and memo_assumptions (nested object of key assumptions).\n"
        "For VENTURE: include ARR, revenue, runway_months, burn_rate, growth_rate.\n"
        "For PE/BUYOUT: include revenue, EBITDA, EBITDA_margin, leverage_at_entry, equity_check, expected_moic, "
        "expected_irr, value_creation_levers, exit_assumptions, management_incentive_structure.\n"
        "For ALL: extract market_size (tam_usd, sam_usd, som_usd) when stated, with tam_description/methodology if available. "
        "Extract red_flags as array of explicit concerns, risks, or concerning language. "
        "For each extracted metric, add to value_explanations with source attribution.\n"
        "Return a single JSON object. Use null when unknown.\n"
        "CURRENCY CONVERSION (always convert to USD before storing any numeric value):\n"
        "- £ (GBP) → multiply by 1.27\n"
        "- € (EUR) → multiply by 1.09\n"
        "- ¥ (JPY) → divide by 154\n"
        "- ₹ (INR) → divide by 84\n"
        "- Note the original currency and amount in value_explanations (e.g. '£2M → $2.54M USD')."
    )
    user_prompt = (
        f"Extract and return a JSON object matching this schema:\n{schema_desc}\n\n"
        f"Document text:\n---\n{text[:120000]}\n---\n\n"
        "Return only the JSON object, no markdown or explanation."
    )
    return system_prompt, user_prompt


def _spreadsheet_prompt(text: str, document_type: str, schema_desc: str, memo_context: Optional[str] = None) -> tuple:
    """Build system and user prompt for spreadsheet extraction (CSV/XLSX management accounts)."""
    system_prompt = (
        "You are a financial analyst. You are given raw spreadsheet data — management accounts, "
        "P&L statements, balance sheets, cash flow statements, or operational dashboards exported as CSV/Excel.\n"
        "Your job is to:\n"
        "1. IDENTIFY the business type: is this a venture-backed startup (ARR/MRR/burn), "
        "a PE-owned company (EBITDA/leverage/FCF), or a traditional business? This determines which metrics matter.\n"
        "2. IDENTIFY the structure: find revenue rows, cost rows, dates/periods in headers or columns.\n"
        "3. EXTRACT the latest-period financial metrics into financial_metrics.\n"
        "   - For ALL companies: revenue, cogs, gross_margin, growth_rate, headcount, cash_balance\n"
        "   - For venture/SaaS: arr, mrr, burn_rate, runway_months, customer_count\n"
        "   - For PE/traditional: ebitda, ebitda_margin, operating_income, net_income, capex, fcf, "
        "total_debt, interest_expense, working_capital, debt_service, tax_expense\n"
        "   Extract WHICHEVER metrics are present — don't force venture metrics on a PE company or vice versa.\n"
        "4. DETECT TRENDS: if multiple periods are present, note growth rates, margin trends, leverage changes.\n"
        "5. ESTIMATE IMPACTS: qualitative signals from the data → impact_estimates.\n"
        "6. For time-series data (monthly/quarterly), extract the full trajectory for forecasting.\n\n"
        "Spreadsheet data is messy — handle merged cells, blank rows, summary rows, and varied date formats.\n"
        "Look for rows labeled: Revenue, ARR, MRR, COGS, Gross Profit, OpEx, EBITDA, Net Income, "
        "Operating Income, Capex, Interest Expense, Debt, Working Capital, Free Cash Flow, "
        "Cash, Burn Rate, Headcount, Customers, Tax, Depreciation, Amortization, etc.\n"
        "When you see monthly columns (Jan, Feb, Mar... or 2024-01, 2024-02...), extract the LATEST non-empty value "
        "and compute growth rates from the series.\n\n"
        "Return a single JSON object. Use null for truly unknown; use empty arrays for missing lists.\n"
        "CURRENCY CONVERSION (always convert to USD before storing any numeric value):\n"
        "- £ (GBP) → multiply by 1.27\n"
        "- € (EUR) → multiply by 1.09\n"
        "- ¥ (JPY) → divide by 154\n"
        "- ₹ (INR) → divide by 84\n"
        "- Note the original currency and amount in value_explanations."
    )
    user_parts = [
        f"Document type: {document_type} (raw spreadsheet data).",
        "",
        "This is tabular financial data. FIRST identify the business type, THEN extract accordingly.",
        "",
        "=== STEP 1: IDENTIFY BUSINESS TYPE ===",
        "Look at the row labels and structure to determine:",
        "  VENTURE/SAAS: Has rows like ARR, MRR, Burn Rate, Runway, CAC, LTV, Customers",
        "  PE/OPERATING: Has rows like EBITDA, Operating Income, Interest Expense, Debt, Capex, D&A, Working Capital",
        "  TRADITIONAL: Has rows like Revenue, COGS, Gross Profit, SG&A, Net Income (no venture or PE-specific rows)",
        "  MULTI-ENTITY: Has divisional/segment columns, intercompany eliminations, consolidated totals",
        "",
        "=== STEP 2: EXTRACT financial_metrics (latest period) ===",
        "  Universal: revenue, cogs, gross_margin, growth_rate, headcount, cash_balance",
        "  Venture/SaaS: arr, mrr, burn_rate, runway_months, customer_count",
        "  PE/Operating/Traditional: ebitda, ebitda_margin, operating_income, net_income, capex, fcf, "
        "total_debt, interest_expense, working_capital, debt_service, tax_expense, depreciation, amortization",
        "  Extract WHICHEVER metrics are present — don't force venture metrics on a PE company or vice versa.",
        "",
        "=== STEP 3: EXTRACT pe_operating_metrics (for PE/operating companies) ===",
        "  If this is a PE portfolio company or traditional operating business, populate pe_operating_metrics:",
        "  - reported_ebitda vs adjusted_ebitda (if add-backs are shown)",
        "  - ebitda_addbacks breakdown (management fees, one-offs, run-rate synergies, non-recurring)",
        "  - Margins: gross_margin_pct, ebitda_margin_pct, operating_margin_pct, net_margin_pct",
        "  - Revenue/EBITDA by segment if divisional reporting exists",
        "  - Leverage: leverage_ratio, interest_coverage, fixed_charge_coverage, dscr",
        "  - Working capital: dso, dpo, dio, cash_conversion_cycle",
        "  - Capex: maintenance_capex vs growth_capex, capex_as_pct_revenue",
        "  - Efficiency: revenue_per_employee, ebitda_per_employee, roic, roe",
        "  - Quality: recurring_revenue_pct, customer_concentration_top5_pct",
        "  - Budget vs actual if present",
        "  - Industry-specific KPIs (utilization, OEE, same-store, fill rate, etc.)",
        "",
        "=== STEP 4: EXTRACT time_series ===",
        "When the spreadsheet has monthly or quarterly columns (Jan, Feb, Mar... or Q1, Q2... or 2024-01, 2024-02...), "
        "extract the FULL TIME SERIES as an array.",
        "",
        "PE/Operating company example:",
        '[{"period": "2025-01", "revenue": 4200000, "cogs": 2100000, "gross_profit": 2100000, '
        '"opex": 1400000, "ebitda": 700000, "operating_income": 550000, '
        '"depreciation": 100000, "amortization": 50000, '
        '"interest_expense": 180000, "tax_expense": 90000, "net_income": 280000, '
        '"capex": 120000, "fcf": 400000, "cash_balance": 3500000, '
        '"total_debt": 12000000, "working_capital": 2800000, '
        '"debt_service": 280000, "dividends": null, "headcount": 185, '
        '"arr": null, "mrr": null, "burn_rate": null, "customers": null}, ...]',
        "",
        "Venture/SaaS example:",
        '[{"period": "2025-01", "revenue": 120000, "arr": 1440000, "mrr": 120000, '
        '"burn_rate": 80000, "cash_balance": 2000000, "headcount": 15, "customers": 45, '
        '"cogs": 25000, "opex": 155000, "gross_profit": 95000, '
        '"ebitda": null, "operating_income": null, "capex": null, "total_debt": null}, ...]',
        "",
        "Preserve EVERY period present in the data. Do not summarize to a single value. "
        "Use null for metrics not present in a given period.",
        "",
        "=== STEP 4b: EXTRACT SUBCATEGORY LINE ITEMS ===",
        "CRITICAL: Real businesses have specific cost line items — extract them ALL as subcategory entries in time_series.",
        "Every row in the spreadsheet that is a specific cost or revenue line item should become a subcategory entry.",
        "The ingestion system handles dynamic subcategories — pass through WHATEVER the business actually reports.",
        "",
        "For each specific line item, add a time_series entry with these fields:",
        '  {"period": "2025-01", "subcategory": "raw_materials", "parent_category": "cogs", "amount": 850000}',
        "",
        "Map each line item to its parent category:",
        "  cogs: direct materials, raw materials, direct labor, factory overhead, freight, packaging, food cost, ",
        "        subcontractors, clinical supplies, third party costs — whatever the business calls its direct costs",
        "  opex_rd: R&D salaries, lab costs, prototyping, clinical trials, product development",
        "  opex_sm: sales salaries, marketing spend, commissions, distribution costs, trade marketing",
        "  opex_ga: rent/occupancy, admin salaries, insurance, legal, accounting, IT, utilities, facility costs",
        "  revenue: subscription revenue, product sales, service revenue, licensing, royalties — by stream if broken out",
        "",
        "Examples by business type:",
        "  Manufacturer: raw_materials, direct_labor, factory_overhead, packaging, freight_out, quality_control, tooling",
        "  Restaurant chain: food_cost, labor, occupancy, supplies, marketing, franchise_fees",
        "  Healthcare: clinical_staff, medical_supplies, lab_costs, pharmacy, insurance_reimbursements, facility_lease",
        "  Distributor: product_purchases, warehousing, freight, fleet_costs, returns_allowance",
        "  Professional services: consultant_salaries, subcontractors, travel, training, utilization_costs",
        "  Software (non-SaaS): license_revenue, maintenance_revenue, professional_services, hosting, support",
        "",
        "Normalize the label: lowercase, underscores instead of spaces (e.g. 'Raw Materials' → 'raw_materials').",
        "Extract EVERY line item — do not roll up into summary categories. The more granular, the better.",
        "",
        "=== STEP 5: ADDITIONAL EXTRACTION ===",
        "- business_updates.latest_update: one-line summary of the financial position",
        "- operational_metrics: headcount, customer_count if present",
        "- impact_estimates: infer from trends (e.g. declining revenue, margin compression, leverage increasing, "
        "working capital deterioration, covenant headroom tightening)",
        "- For PE companies: attempt to fill ebitda_bridge and thesis_tracking in impact_estimates if enough data exists",
        "- value_explanations: cite the specific cells/values you used",
        "",
        "Schema (JSON):",
        schema_desc,
        "",
        "Spreadsheet data:",
        "---",
        (text[:120000] if len(text) > 120000 else text),
        "---",
        "Return only the JSON object, no markdown or explanation.",
    ]
    if memo_context:
        user_parts.insert(1, f"BASELINE ANCHOR (same company, use for scaling all estimates): {memo_context}")
    user_prompt = "\n".join(user_parts)
    return system_prompt, user_prompt


def _flat_prompt(text: str, document_type: str, schema_desc: str) -> tuple:
    """Build system and user prompt for flat schema (pitch_deck / other)."""
    system_prompt = (
        "You are a financial document analyst. Extract structured data from the given document text. "
        "This could be a venture pitch deck, a PE CIM (confidential information memorandum), a company overview, or any business document.\n"
        "FIRST identify the business type:\n"
        "  - VENTURE/SAAS: ARR, MRR, burn, runway, funding rounds → extract arr, burn_rate, stage, total_funding\n"
        "  - PE/OPERATING/TRADITIONAL: revenue, EBITDA, margins, leverage, FCF → extract revenue, ebitda, ebitda_margin, enterprise_value\n"
        "Return a single JSON object with these keys (use null when unknown): "
        "company_name, revenue, arr, ebitda, ebitda_margin, stage, total_funding, valuation, enterprise_value, "
        "leverage_ratio, key_metrics (array of strings), summary, sector, target_market, business_model. "
        "Numbers must be numeric (no currency symbols in values). key_metrics is an array of short strings — "
        "for PE companies include metrics like 'EBITDA margin 22%', 'leverage 3.5x', 'FCF yield 8%'. "
        "For each extracted metric, add to value_explanations with source attribution.\n"
        "CURRENCY CONVERSION (always convert to USD before storing any numeric value):\n"
        "- £ (GBP) → multiply by 1.27\n"
        "- € (EUR) → multiply by 1.09\n"
        "- ¥ (JPY) → divide by 154\n"
        "- ₹ (INR) → divide by 84\n"
        "- Note the original currency and amount in value_explanations (e.g. '£2M → $2.54M USD')."
    )
    user_prompt = (
        f"Document type: {document_type}\n\n"
        f"Extract and return a JSON object matching this schema:\n{schema_desc}\n\n"
        f"Document text:\n---\n{text[:120000]}\n---\n\n"
        "Return only the JSON object, no markdown or explanation."
    )
    return system_prompt, user_prompt


async def _extract_document_structured_async(
    text: str,
    document_type: str,
    memo_context: Optional[str] = None,
    erp_category_hint: Optional[str] = None,
    erp_subcategory_hint: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Call model_router with a prompt and JSON schema to extract structured data from document text.
    Branches by document_type:
    - legal doc types → legal clause extraction schema
    - monthly_update/board_deck → signal schema
    - investment_memo → memo schema
    - else → flat
    Returns a dict suitable for extracted_data (normalized so financial_metrics and period_date exist where applicable).
    """
    from app.services.model_router import ModelRouter, ModelCapability

    # Create a fresh router instance — NOT the singleton.
    # This function runs inside asyncio.run() in a worker thread, which creates
    # a new event loop. The singleton's async clients (AsyncAnthropic, aiohttp)
    # are bound to the main event loop and will deadlock here. A fresh instance
    # initializes its clients on this thread's loop.
    router = ModelRouter()
    doc_type = (document_type or "other").strip().lower()

    if doc_type in LEGAL_DOC_TYPES:
        # Legal document — clause extraction with parent-child hierarchy
        legal_schema = _get_legal_schema(doc_type)
        schema_desc = json.dumps(legal_schema, indent=2)
        system_prompt, user_prompt = _legal_extraction_prompt(
            text, doc_type, schema_desc,
            erp_category_hint=erp_category_hint,
            erp_subcategory_hint=erp_subcategory_hint,
        )
        empty = _empty_legal_extraction()
    elif doc_type == "investment_memo":
        schema_desc = json.dumps(INVESTMENT_MEMO_SCHEMA, indent=2)
        system_prompt, user_prompt = _memo_prompt(text, schema_desc)
        empty = _empty_memo_extraction()
    elif doc_type == "financial_statement":
        # Spreadsheet data (CSV/XLSX) — use spreadsheet-specific prompt that handles
        # tabular management accounts, P&Ls, etc. Same signal schema, different prompt.
        schema_desc = json.dumps(COMPANY_UPDATE_SIGNAL_SCHEMA, indent=2)
        system_prompt, user_prompt = _spreadsheet_prompt(text, document_type, schema_desc, memo_context)
        empty = _empty_signal_extraction()
    else:
        # ALL non-memo docs use signal-first extraction — extracts business_updates,
        # operational_metrics, impact_estimates, financial_metrics, etc.
        schema_desc = json.dumps(COMPANY_UPDATE_SIGNAL_SCHEMA, indent=2)
        system_prompt, user_prompt = _signal_first_prompt(text, document_type, schema_desc, memo_context)
        empty = _empty_signal_extraction()

    # Legal docs get more tokens — clause extraction is verbose.
    # Use 16384 for legal to avoid truncation on complex contracts.
    max_tok = 16384 if doc_type in LEGAL_DOC_TYPES else 4096

    # Legal docs MUST use a quality model — Haiku truncates at 8k tokens
    # and produces incomplete JSON.  Explicit preferred_models overrides
    # the caller_context routing so we guarantee Sonnet/GPT-5.2 first.
    legal_preferred = ["gpt-5.2", "claude-sonnet-4-6", "gemini-2.5-pro"] if doc_type in LEGAL_DOC_TYPES else None

    try:
        result = await router.get_completion(
            prompt=user_prompt,
            system_prompt=system_prompt,
            capability=ModelCapability.STRUCTURED,
            max_tokens=max_tok,
            temperature=0.2,
            json_mode=True,
            preferred_models=legal_preferred,
            caller_context="document_process_service.extract_structured",
        )
        raw = (result.get("response") or "").strip()
        if not raw:
            return empty

        parsed = _extract_json_object(raw)
        if isinstance(parsed, dict):
            # Legal docs return their own shape — skip financial normalization
            if doc_type in LEGAL_DOC_TYPES:
                return _normalize_legal_extraction(parsed)
            return _normalize_extraction(parsed, document_type=document_type)
        return empty
    except Exception as e:
        logger.exception("extract_document_structured failed: %s", e)
        if doc_type in LEGAL_DOC_TYPES:
            out = _empty_legal_extraction()
        elif doc_type == "investment_memo":
            out = _empty_memo_extraction()
        else:
            out = _empty_signal_extraction()
        out["_extraction_error"] = str(e)
        return out


def _extract_json_object(raw: str) -> dict:
    """Robustly extract a JSON object from an LLM response.

    Handles preamble text, code fences, and the ``[`` prefill artefact
    that can cause Claude to emit ``[PROCESSING DOCUMENT...]`` before the
    actual JSON payload.
    """
    # 1. Strip code fences anywhere in the string
    cleaned = re.sub(r"```(?:json)?", "", raw).strip()

    # 2. Try direct parse first (fast path)
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, dict):
            return obj
    except json.JSONDecodeError:
        pass

    # 3. Locate the first '{' and try to parse from there
    start = cleaned.find("{")
    if start != -1:
        # Find matching closing brace by trying successively shorter slices
        depth = 0
        end = None
        for i in range(start, len(cleaned)):
            if cleaned[i] == "{":
                depth += 1
            elif cleaned[i] == "}":
                depth -= 1
                if depth == 0:
                    end = i + 1
                    break
        if end is not None:
            try:
                return json.loads(cleaned[start:end])
            except json.JSONDecodeError:
                pass

    raise json.JSONDecodeError("No JSON object found in response", raw, 0)


def _empty_extraction(error: Optional[str] = None) -> Dict[str, Any]:
    out: Dict[str, Any] = {
        "company_name": None,
        "revenue": None,
        "arr": None,
        "stage": None,
        "total_funding": None,
        "valuation": None,
        "key_metrics": [],
        "summary": "",
        "sector": None,
        "target_market": None,
        "business_model": None,
        "value_explanations": {},
    }
    if error:
        out["_extraction_error"] = error
    return out


def _empty_signal_extraction(error: Optional[str] = None) -> Dict[str, Any]:
    out: Dict[str, Any] = {
        "company_name": None,
        "summary": "",
        "business_updates": {
            "product_updates": [],
            "achievements": [],
            "challenges": [],
            "risks": [],
            "key_milestones": [],
            "asks": [],
            "latest_update": "",
            "defensive_language": [],
        },
        "operational_metrics": {
            "new_hires": [],
            "headcount": None,
            "customer_count": None,
            "enterprise_customers": None,
            "smb_customers": None,
        },
        "extracted_entities": {
            "competitors_mentioned": [],
            "industry_terms": [],
            "partners_mentioned": [],
        },
        "red_flags": [],
        "implications": [],
        "period_date": None,
        "financial_metrics": {},
        "value_explanations": {},
    }
    if error:
        out["_extraction_error"] = error
    return out


def _empty_memo_extraction(error: Optional[str] = None) -> Dict[str, Any]:
    out: Dict[str, Any] = {
        "company_name": None,
        "investment_date": None,
        "round": None,
        "valuation_pre_money": None,
        "deal_terms_summary": None,
        "memo_assumptions": {},
        "revenue": None,
        "arr": None,
        "runway_months": None,
        "stage": None,
        "total_funding": None,
        "valuation": None,
        "key_metrics": [],
        "summary": "",
        "sector": None,
        "target_market": None,
        "business_model": None,
        "market_size": None,
        "red_flags": [],
        "value_explanations": {},
    }
    if error:
        out["_extraction_error"] = error
    return out


def _ensure_numeric(val: Any) -> Optional[float]:
    if val is None:
        return None
    if isinstance(val, (int, float)) and not isinstance(val, bool):
        return float(val)
    try:
        s = str(val).replace(",", "").replace("$", "").strip()
        return float(s) if s else None
    except (ValueError, TypeError):
        return None


def _normalize_extraction(d: Dict[str, Any], document_type: Optional[str] = None) -> Dict[str, Any]:
    """
    Normalize extraction output for DB: ensure financial_metrics and period_date exist;
    map flat revenue/arr into financial_metrics for backward compatibility.
    """
    doc_type = (document_type or "other").strip().lower()

    # —— Signal shape (monthly_update / board_deck / board_transcript) ——
    if doc_type != "investment_memo":
        out = dict(d)
        bu = out.get("business_updates")
        if not isinstance(bu, dict):
            out["business_updates"] = (_empty_signal_extraction().get("business_updates") or {}).copy()
        else:
            for key in ("product_updates", "achievements", "challenges", "risks", "key_milestones", "asks", "defensive_language"):
                if key in out["business_updates"] and not isinstance(out["business_updates"][key], list):
                    out["business_updates"][key] = []
            if "latest_update" not in out["business_updates"]:
                out["business_updates"]["latest_update"] = ""
        om = out.get("operational_metrics")
        if not isinstance(om, dict):
            out["operational_metrics"] = (_empty_signal_extraction().get("operational_metrics") or {}).copy()
        ee = out.get("extracted_entities")
        if not isinstance(ee, dict):
            out["extracted_entities"] = (_empty_signal_extraction().get("extracted_entities") or {}).copy()
        fm = out.get("financial_metrics")
        if not isinstance(fm, dict):
            fm = {}
        out["financial_metrics"] = {
            "arr": _ensure_numeric(fm.get("arr")),
            "revenue": _ensure_numeric(fm.get("revenue")),
            "mrr": _ensure_numeric(fm.get("mrr")),
            "burn_rate": _ensure_numeric(fm.get("burn_rate")),
            "runway_months": _ensure_numeric(fm.get("runway_months")),
            "cash_balance": _ensure_numeric(fm.get("cash_balance")),
            "gross_margin": _ensure_numeric(fm.get("gross_margin")),
            "growth_rate": _ensure_numeric(fm.get("growth_rate")),
            "customer_count": _ensure_numeric(fm.get("customer_count")),
        }

        out["period_date"] = out.get("period_date") if isinstance(out.get("period_date"), str) else None
        # Preserve impact_estimates from LLM extraction (transformation layer)
        ie = out.get("impact_estimates")
        if isinstance(ie, dict):
            out["impact_estimates"] = {
                "estimated_arr_impact": _ensure_numeric(ie.get("estimated_arr_impact")),
                "estimated_burn_impact": _ensure_numeric(ie.get("estimated_burn_impact")),
                "estimated_runway_impact": _ensure_numeric(ie.get("estimated_runway_impact")),
                "estimated_headcount_impact": _ensure_numeric(ie.get("estimated_headcount_impact")),
                "estimated_cash_impact": _ensure_numeric(ie.get("estimated_cash_impact")),
                "estimated_valuation_impact": _ensure_numeric(ie.get("estimated_valuation_impact")),
                "estimated_growth_rate_change": _ensure_numeric(ie.get("estimated_growth_rate_change")),
                "impact_reasoning": ie.get("impact_reasoning") if isinstance(ie.get("impact_reasoning"), dict) else {},
            }
        else:
            out["impact_estimates"] = None
        # Unified shape for suggestions: add company_info, growth_metrics, runway_and_cash
        out["company_info"] = {
            "name": out.get("company_name"),
            "sector": out.get("sector"),
            "stage": out.get("stage"),
            "valuation": out.get("valuation"),
            "funding_raised": out.get("total_funding") or out.get("total_raised"),
            "industry": out.get("sector"),
            "business_model": out.get("business_model"),
            "category": out.get("category"),
        }
        # growth_rate is canonical decimal (0.3 = 30%). Convert to percentage for display field.
        g = _ensure_numeric(out["financial_metrics"].get("growth_rate"))
        growth_annual = (g * 100) if g is not None else None
        out["growth_metrics"] = {
            "current_arr": out["financial_metrics"].get("arr"),
            "revenue_growth_annual_pct": growth_annual,
            "revenue_growth_monthly_pct": None,
        }
        out["runway_and_cash"] = {
            "runway_months": out["financial_metrics"].get("runway_months"),
            "cash_in_bank": out["financial_metrics"].get("cash_balance"),
            "burn_rate": out["financial_metrics"].get("burn_rate"),
        }
        # Canonical: market_size, red_flags, implications for suggestions/analysis
        out["red_flags"] = [x for x in (out.get("red_flags") or []) if isinstance(x, str)]
        out["implications"] = [x for x in (out.get("implications") or []) if isinstance(x, str)]
        ms = out.get("market_size")
        if isinstance(ms, dict) and any(ms.get(k) is not None for k in ("tam_usd", "sam_usd", "som_usd")):
            out["market_size"] = {
                "tam_usd": _ensure_numeric(ms.get("tam_usd")),
                "sam_usd": _ensure_numeric(ms.get("sam_usd")),
                "som_usd": _ensure_numeric(ms.get("som_usd")),
                "tam_description": ms.get("tam_description") if isinstance(ms.get("tam_description"), str) else None,
                "methodology": ms.get("methodology") if isinstance(ms.get("methodology"), str) else None,
            }
        else:
            out["market_size"] = None
        ve = out.get("value_explanations") if isinstance(out.get("value_explanations"), dict) else {}
        out["value_explanations"] = ve
        if ve:
            logger.debug("value_explanations (signal): %s", list(ve.keys()))
        return out

    # —— Memo shape ——
    if doc_type == "investment_memo":
        keys = set(INVESTMENT_MEMO_SCHEMA.keys())
        out = {}
        for k in keys:
            v = d.get(k)
            if k == "key_metrics" and not isinstance(v, list):
                out[k] = [str(x) for x in (v or [])] if isinstance(v, (list, tuple)) else []
            elif k == "memo_assumptions":
                out[k] = v if isinstance(v, dict) else {}
            elif k in ("revenue", "arr", "total_funding", "valuation", "valuation_pre_money", "runway_months") and v is not None:
                out[k] = _ensure_numeric(v)
            else:
                out[k] = v if isinstance(v, (str, int, float, list, dict, type(None))) else (str(v) if v is not None else None)
        fm = {
            "arr": out.get("arr") or _ensure_numeric(d.get("arr")),
            "revenue": out.get("revenue") or _ensure_numeric(d.get("revenue")),
            "runway_months": out.get("runway_months") or _ensure_numeric(d.get("runway_months")),
            "burn_rate": None,
            "cash_balance": None,
            "gross_margin": None,
            "growth_rate": None,
        }
        out["financial_metrics"] = fm
        out["period_date"] = out.get("investment_date") or _iso_now()[:10]
        out["company_info"] = {
            "name": out.get("company_name"),
            "sector": out.get("sector"),
            "stage": out.get("stage"),
            "valuation": out.get("valuation") or out.get("valuation_pre_money"),
            "funding_raised": out.get("total_funding"),
            "industry": out.get("sector"),
        }
        out["growth_metrics"] = {
            "current_arr": out.get("arr"),
            "revenue_growth_annual_pct": None,
            "revenue_growth_monthly_pct": None,
        }
        out["business_updates"] = {
            "latest_update": (out.get("summary") or "")[:2000] if out.get("summary") else "",
            "product_updates": out.get("key_metrics") if isinstance(out.get("key_metrics"), list) else [],
            "achievements": [],
            "challenges": [],
            "risks": [],
            "key_milestones": [],
            "asks": [],
            "defensive_language": [],
        }
        out["runway_and_cash"] = {
            "runway_months": out.get("runway_months") or fm.get("runway_months"),
            "cash_in_bank": None,
            "burn_rate": None,
        }
        # Canonical: market_size, red_flags for suggestions/analysis
        out["red_flags"] = [x for x in (out.get("red_flags") or d.get("red_flags") or []) if isinstance(x, str)]
        ms = out.get("market_size") or d.get("market_size")
        if isinstance(ms, dict) and any(ms.get(k) is not None for k in ("tam_usd", "sam_usd", "som_usd")):
            out["market_size"] = {
                "tam_usd": _ensure_numeric(ms.get("tam_usd")),
                "sam_usd": _ensure_numeric(ms.get("sam_usd")),
                "som_usd": _ensure_numeric(ms.get("som_usd")),
                "tam_description": ms.get("tam_description") if isinstance(ms.get("tam_description"), str) else None,
                "methodology": ms.get("methodology") if isinstance(ms.get("methodology"), str) else None,
            }
        else:
            out["market_size"] = None
        ve = out.get("value_explanations") if isinstance(out.get("value_explanations"), dict) else {}
        out["value_explanations"] = ve
        if ve:
            logger.debug("value_explanations (memo): %s", list(ve.keys()))
        return out

    # —— Flat shape (other doc types) ——
    keys = set(DOCUMENT_EXTRACTION_SCHEMA.keys())
    out = {}
    for k in keys:
        v = d.get(k)
        if k == "key_metrics" and not isinstance(v, list):
            out[k] = [str(x) for x in (v or [])] if isinstance(v, (list, tuple)) else []
        elif k == "red_flags" and isinstance(v, list):
            out[k] = [x for x in v if isinstance(x, str)]
        elif k == "market_size" and isinstance(v, dict):
            out[k] = v  # normalized below
        elif k in ("revenue", "arr", "total_funding", "valuation") and v is not None:
            out[k] = _ensure_numeric(v)
        else:
            out[k] = v if isinstance(v, (str, int, float, list, dict, type(None))) else (str(v) if v is not None else None)
    out["financial_metrics"] = {
        "arr": out.get("arr"),
        "revenue": out.get("revenue"),
        "mrr": None,
        "burn_rate": None,
        "runway_months": None,
        "cash_balance": None,
        "gross_margin": None,
        "growth_rate": None,
    }
    out["period_date"] = _iso_now()[:10]
    # Unified shape for suggestions/analysis: company_info and growth_metrics
    out["company_info"] = {
        "name": out.get("company_name"),
        "sector": out.get("sector"),
        "stage": out.get("stage"),
        "valuation": out.get("valuation"),
        "funding_raised": out.get("total_funding"),
        "industry": out.get("sector"),
    }
    out["growth_metrics"] = {
        "current_arr": out.get("arr"),
        "revenue_growth_annual_pct": None,
        "revenue_growth_monthly_pct": None,
    }
    # business_updates placeholder so suggestions can read latest_update from summary
    out["business_updates"] = {
        "latest_update": (out.get("summary") or "")[:2000] if out.get("summary") else "",
        "product_updates": out.get("key_metrics") if isinstance(out.get("key_metrics"), list) else [],
        "achievements": [],
        "challenges": [],
        "risks": [],
        "key_milestones": [],
        "asks": [],
        "defensive_language": [],
    }
    out["runway_and_cash"] = {"runway_months": None, "cash_in_bank": None, "burn_rate": None}
    out["operational_metrics"] = {"new_hires": [], "headcount": None, "customer_count": None}
    # Canonical: market_size, red_flags
    out["red_flags"] = [x for x in (out.get("red_flags") or d.get("red_flags") or []) if isinstance(x, str)]
    ms = out.get("market_size") or d.get("market_size")
    if isinstance(ms, dict) and any(ms.get(k) is not None for k in ("tam_usd", "sam_usd", "som_usd")):
        out["market_size"] = {
            "tam_usd": _ensure_numeric(ms.get("tam_usd")),
            "sam_usd": _ensure_numeric(ms.get("sam_usd")),
            "som_usd": _ensure_numeric(ms.get("som_usd")),
            "tam_description": ms.get("tam_description") if isinstance(ms.get("tam_description"), str) else None,
            "methodology": ms.get("methodology") if isinstance(ms.get("methodology"), str) else None,
        }
    else:
        out["market_size"] = None
    ve = out.get("value_explanations") if isinstance(out.get("value_explanations"), dict) else {}
    out["value_explanations"] = ve
    if ve:
        logger.debug("value_explanations (flat): %s", list(ve.keys()))
    return out


# ---------------------------------------------------------------------------
# Direct upsert of extracted clauses → document_clauses (structured columns)
# ---------------------------------------------------------------------------

def _enrich_clause_flags(
    clause: Dict[str, Any],
    extracted_data: Dict[str, Any],
) -> list:
    """Dynamically detect risk flags from extracted clause data.

    Goes beyond the template flags the LLM returns (above_market, auto_renew_risk)
    by actually inspecting the data for concrete risks.
    """
    flags = list(clause.get("flags") or [])
    clause_type = (clause.get("clause_type") or "").lower()
    obligations = clause.get("obligations") or []
    xrefs = clause.get("cross_references") or []
    key_dates = extracted_data.get("key_dates") or []

    # 1. Deadline proximity — obligation due within 90 days (concrete dates, not generic labels)
    from datetime import datetime as _dt, timedelta as _td
    now = _dt.utcnow()
    for ob in obligations:
        deadline_str = ob.get("deadline")
        if not deadline_str:
            continue
        try:
            deadline = _dt.fromisoformat(str(deadline_str).replace("Z", "+00:00").replace("+00:00", ""))
        except (ValueError, TypeError):
            continue
        days_out = (deadline - now).days
        date_fmt = deadline.strftime("%d %b %Y")
        if 0 < days_out <= 90:
            desc = ob.get("description", "obligation")[:60]
            flags.append(f"Deadline in {days_out} days ({date_fmt}) — {desc}")
        elif days_out <= 0:
            desc = ob.get("description", "obligation")[:60]
            flags.append(f"Deadline passed ({date_fmt}) — {desc}")

    # 2. Auto-renewal trap — auto_renewal clause without termination_for_convenience
    if clause_type in ("auto_renewal", "renewal"):
        has_convenience_termination = any(
            c.get("clause_type") == "termination" and "convenience" in (c.get("text") or "").lower()
            for c in (extracted_data.get("clauses") or [])
        )
        if not has_convenience_termination:
            flags.append("Auto-renews with no termination for convenience")

    # 3. Exposure detection — uncapped liability, personal guarantees, cross-defaults
    if clause_type == "liability_cap":
        text = (clause.get("text") or "").lower()
        if "unlimited" in text or "uncapped" in text or "no limit" in text:
            flags.append("Uncapped liability — unlimited exposure")
    if clause_type in ("personal_guarantee", "parent_guarantee"):
        flags.append("Personal guarantee — individual liability for principals")
    if clause_type == "cross_default":
        flags.append("Cross-default — breach here triggers default on other agreements")

    # 4. Key dates approaching — from document-level key_dates
    for kd in key_dates:
        kd_date_str = kd.get("date")
        auto_action = kd.get("auto_action")
        if not kd_date_str:
            continue
        try:
            kd_date = _dt.fromisoformat(str(kd_date_str).replace("Z", "+00:00").replace("+00:00", ""))
        except (ValueError, TypeError):
            continue
        days_out = (kd_date - now).days
        date_fmt = kd_date.strftime("%d %b %Y")
        event = kd.get("event", auto_action or "event")[:40]
        if 0 < days_out <= 60 and auto_action == "auto_renew":
            flags.append(f"Auto-renews in {days_out} days ({date_fmt}) — {event}")
        elif 0 < days_out <= 30 and auto_action == "terminate":
            flags.append(f"Termination in {days_out} days ({date_fmt}) — {event}")

    # 5. High-value cross-references — flags clauses that define financial engine inputs
    for xref in xrefs:
        relationship = (xref.get("relationship") or "").lower()
        field = xref.get("field", "")
        to_service = xref.get("to_service", "")
        if relationship in ("defines", "overrides"):
            flags.append(f"{relationship.title()}s {to_service}.{field}" if field else f"{relationship.title()}s {to_service}")
            break

    return flags


def _upsert_clauses_to_document_clauses(
    *,
    extracted_data: Dict[str, Any],
    document_id: str,
    document_name: str,
    company_id: Optional[str],
    fund_id: Optional[str],
    document_type: str,
) -> int:
    """Upsert extracted clauses directly into document_clauses with structured columns.

    This is the real fix: obligations, deadlines, amounts, flags go into actual
    SQL columns — not into a JSONB blob in pending_suggestions.
    The orchestrator, legal grid, debt/legal context queries all read from
    document_clauses, so this makes extracted data immediately queryable.
    """
    clauses = extracted_data.get("clauses")
    if not isinstance(clauses, list) or not clauses:
        return 0

    erp = extracted_data.get("erp_attribution") or {}
    parties = extracted_data.get("parties") or []
    value_explanations = extracted_data.get("value_explanations") or {}
    red_flags = extracted_data.get("red_flags") or []

    from app.core.supabase_client import get_supabase_client
    sb = get_supabase_client()
    if not sb:
        logger.warning("[CLAUSE_UPSERT] Supabase unavailable — clauses not persisted")
        return 0

    # Document-level fields for grid columns (not per-clause in the DB schema,
    # so we store in metadata and let the frontend route read them)
    doc_effective_date = extracted_data.get("effective_date")
    doc_expiry_date = extracted_data.get("expiration_date") or extracted_data.get("expiry_date")
    doc_total_value = (
        _ensure_numeric(extracted_data.get("investment_amount"))
        or _ensure_numeric(erp.get("annual_value"))
        or _ensure_numeric(extracted_data.get("valuation_post_money"))
    )
    # Counterparty: second party in parties list (first is usually "us")
    doc_counterparty = ""
    if len(parties) >= 2:
        doc_counterparty = parties[1].get("name", "")
    elif len(parties) == 1:
        doc_counterparty = parties[0].get("name", "")

    rows = []
    for clause in clauses:
        if not isinstance(clause, dict):
            continue
        clause_id = clause.get("id", "")
        if not clause_id:
            continue

        # Enrich flags with dynamic detection
        enriched_flags = _enrich_clause_flags(clause, extracted_data)

        # Primary obligation
        obligations = clause.get("obligations") or []
        primary_ob = obligations[0] if obligations else {}

        # Primary cross-reference
        xrefs = clause.get("cross_references") or []
        primary_xref = xrefs[0] if xrefs else {}

        # Party from obligation or document-level
        party = primary_ob.get("party", "")
        if not party and parties:
            party = parties[0].get("name", "")

        # Parse obligation_deadline as a date string (YYYY-MM-DD) or None
        deadline_raw = primary_ob.get("deadline")
        deadline_date = None
        if deadline_raw:
            try:
                # Accept ISO formats
                from datetime import datetime as _dt
                parsed = _dt.fromisoformat(str(deadline_raw).replace("Z", "+00:00").replace("+00:00", ""))
                deadline_date = parsed.strftime("%Y-%m-%d")
            except (ValueError, TypeError):
                # If it's already YYYY-MM-DD, keep it
                if re.match(r"^\d{4}-\d{2}-\d{2}$", str(deadline_raw)):
                    deadline_date = str(deadline_raw)

        # Build reasoning: prefer LLM's value_explanation (actual analysis),
        # fall back to flags. Never just copy raw clause text — that's useless.
        clause_id_str = clause.get("id", "")
        explanation = (
            value_explanations.get(clause_id_str)
            or value_explanations.get(clause.get("clause_type", ""))
            or ""
        )
        if explanation:
            # LLM wrote a real explanation — use it directly
            reasoning = explanation[:500]
        elif enriched_flags:
            # No explanation but we have flags — use those as reasoning
            reasoning = " · ".join(enriched_flags)
        else:
            # Last resort — clause title/type, not raw text
            title = clause.get("title") or clause.get("clause_type") or "clause"
            reasoning = f"Extracted: {title}"

        # Confidence — clauses with flags get higher confidence (flags are now free-form)
        confidence = 0.80
        if enriched_flags:
            # Any flagged clause is higher confidence — the LLM found something notable
            # Risk-indicating keywords in free-form flags boost further
            flag_text = " ".join(enriched_flags).lower()
            risk_keywords = ("uncapped", "unlimited", "no cap", "personal", "guarantee",
                             "cross-default", "passed", "aggressive", "above market",
                             "no termination", "locked in", "one-sided")
            if any(kw in flag_text for kw in risk_keywords):
                confidence = 0.92
            else:
                confidence = 0.85

        row = {
            "fund_id": fund_id or _UNLINKED,
            "company_id": company_id or _UNLINKED,
            "document_id": document_id,
            "document_name": document_name or f"doc:{document_id}",
            "clause_id": clause_id,
            "title": clause.get("title", ""),
            "clause_type": clause.get("clause_type", "other"),
            "clause_text": clause_text,
            "party": party,
            "flags": enriched_flags,  # TEXT[] — real array, not comma-separated
            "obligation_desc": primary_ob.get("description", "") or "",
            "obligation_deadline": deadline_date,
            "cross_ref_service": primary_xref.get("to_service", "") or "",
            "cross_ref_field": primary_xref.get("field", "") or "",
            "cross_ref_value": str(primary_xref.get("value", "")) if primary_xref.get("value") is not None else "",
            "erp_category": erp.get("category") or "",
            "erp_subcategory": erp.get("subcategory") or "",
            "annual_value": _ensure_numeric(erp.get("annual_value")),
            "monthly_amount": _ensure_numeric(erp.get("monthly_amount")),
            "reasoning": reasoning,
            "confidence": confidence,
            "source_service": f"document:{document_id}",
            "metadata": {
                "document_type": document_type,
                "parent_clause_id": clause.get("parent_id"),
                "children_clause_ids": clause.get("children", []),
                "all_obligations": obligations,
                "all_cross_references": xrefs,
                "red_flags": red_flags,
                "enriched_flags": [f for f in enriched_flags if f not in (clause.get("flags") or [])],
                # Document-level fields for grid columns (no dedicated DB columns)
                "effective_date": doc_effective_date,
                "expiry_date": doc_expiry_date,
                "total_value": doc_total_value,
                "counterparty": doc_counterparty,
            },
        }
        rows.append(row)

    if not rows:
        return 0

    try:
        sb.table("document_clauses").upsert(
            rows,
            on_conflict="fund_id,company_id,clause_id,document_id",
        ).execute()
        logger.info(
            "[CLAUSE_UPSERT] Wrote %d clauses to document_clauses for doc %s (type=%s)",
            len(rows), document_id, document_type,
        )
        return len(rows)
    except Exception as e:
        logger.warning("[CLAUSE_UPSERT] Batch upsert failed for %s: %s — falling back to individual", document_id, e)
        written = 0
        for row in rows:
            try:
                sb.table("document_clauses").upsert(
                    row,
                    on_conflict="fund_id,company_id,clause_id,document_id",
                ).execute()
                written += 1
            except Exception as e2:
                logger.warning("[CLAUSE_UPSERT] Individual write failed for %s.%s: %s", document_id, row["clause_id"], e2)
        return written


def run_post_extraction_pipeline(
    *,
    extracted_data: Dict[str, Any],
    document_id: str,
    document_type: str,
    company_id: Optional[str],
    fund_id: Optional[str],
    document_name: str,
    field_count: int = 0,
) -> None:
    """Run the full post-extraction pipeline: suggestions, bridges, cap table, actuals.

    Called by both run_document_process (single-doc) and the parallel batch path.
    All errors are caught and logged — never raises.
    """
    doc_type_for_routing = (document_type or "").strip().lower()

    # ── Portfolio suggestions (non-legal, non-financial-statement) ──
    if extracted_data and doc_type_for_routing not in LEGAL_DOC_TYPES and doc_type_for_routing != "financial_statement":
        try:
            from app.services.micro_skills.suggestion_emitter import emit_document_suggestions
            n = emit_document_suggestions(
                extracted_data=extracted_data,
                company_id=company_id,
                fund_id=fund_id,
                document_id=document_id,
                document_name=document_name,
            )
            if n:
                logger.info("Emitted %d portfolio suggestions from document %s (type=%s)", n, document_id, doc_type_for_routing)
            else:
                logger.info(
                    "[DOC_PROCESS] No portfolio suggestions emitted from document %s "
                    "(extracted %d fields, company_id=%s, fund_id=%s)",
                    document_id, field_count, company_id, fund_id,
                )
        except Exception as e:
            logger.warning("Failed to emit document suggestions for %s: %s", document_id, e, exc_info=True)
    elif doc_type_for_routing in LEGAL_DOC_TYPES:
        logger.info("[DOC_PROCESS] Skipping portfolio emitter for legal doc %s (type=%s) — legal emitter handles this", document_id, doc_type_for_routing)
    elif doc_type_for_routing == "financial_statement":
        logger.info("[DOC_PROCESS] Skipping portfolio emitter for financial_statement doc %s — actuals ingestion handles this", document_id)

    # ── Legal clause suggestions ──
    if extracted_data and doc_type_for_routing in LEGAL_DOC_TYPES:
        try:
            from app.services.micro_skills.suggestion_emitter import emit_legal_suggestions
            n = emit_legal_suggestions(
                extracted_data=extracted_data,
                company_id=company_id,
                fund_id=fund_id,
                document_id=document_id,
                document_name=document_name,
            )
            if n:
                logger.info("Emitted %d legal clause suggestions from document %s", n, document_id)
        except Exception as e:
            logger.warning("Failed to emit legal suggestions for %s: %s", document_id, e, exc_info=True)

        # NOTE: Direct upsert to document_clauses removed — suggestions-first flow.
        # Clauses go through pending_suggestions (via emit_legal_suggestions above).
        # User reviews and accepts; acceptance handler in frontend upserts to document_clauses.

    # ── Contract → P&L bridge (ERP attribution) ──
    if extracted_data and extracted_data.get("erp_attribution"):
        from app.services.contract_pnl_bridge import COMMERCIAL_DOC_TYPES, INTERCOMPANY_DOC_TYPES
        if doc_type_for_routing in COMMERCIAL_DOC_TYPES or doc_type_for_routing in LEGAL_DOC_TYPES:
            try:
                from app.services.contract_pnl_bridge import bridge_contract_to_pnl
                pnl_result = bridge_contract_to_pnl(
                    extracted_data=extracted_data,
                    company_id=company_id,
                    fund_id=fund_id,
                    document_id=document_id,
                    document_type=doc_type_for_routing,
                    document_name=document_name,
                )
                if pnl_result.get("success"):
                    logger.info(
                        "[DOC_PNL_BRIDGE] Wrote %d P&L rows for doc %s (%s)",
                        pnl_result["rows_written"], document_id,
                        pnl_result.get("details", {}).get("hierarchy_path", ""),
                    )
            except Exception as e:
                logger.warning("[DOC_PNL_BRIDGE] Failed for %s: %s", document_id, e, exc_info=True)

        # Intercompany agreements → TP engine suggestions
        if doc_type_for_routing in INTERCOMPANY_DOC_TYPES:
            try:
                from app.services.contract_pnl_bridge import bridge_contract_to_tp
                tp_result = bridge_contract_to_tp(
                    extracted_data=extracted_data,
                    company_id=company_id,
                    document_id=document_id,
                    document_name=document_name,
                )
                if tp_result.get("success"):
                    logger.info(
                        "[DOC_TP_BRIDGE] Created IC suggestion for doc %s (%s)",
                        document_id, tp_result.get("transaction_type"),
                    )
            except Exception as e:
                logger.warning("[DOC_TP_BRIDGE] Failed for %s: %s", document_id, e, exc_info=True)

    # ── Time-series actuals ──
    if extracted_data and extracted_data.get("time_series"):
        try:
            from app.services.actuals_ingestion import ingest_time_series
            n = ingest_time_series(
                time_series=extracted_data["time_series"],
                company_id=company_id,
                fund_id=fund_id,
                document_id=document_id,
            )
            if n:
                logger.info("Ingested %d actuals rows from document %s", n, document_id)
        except Exception as e:
            logger.warning("Failed to ingest actuals for %s: %s", document_id, e, exc_info=True)

    # ── Path A: Document-derived cap table (SHA, term_sheet, etc. with cap table xrefs) ──
    cap_table_doc_types = {"sha", "term_sheet", "side_letter", "option_agreement", "spa", "convertible_note", "safe"}
    if extracted_data and doc_type_for_routing in cap_table_doc_types:
        has_cap_table_xrefs = any(
            xr.get("to_service") == "cap_table"
            for clause in (extracted_data.get("clauses") or [])
            if isinstance(clause, dict)
            for xr in (clause.get("cross_references") or [])
            if isinstance(xr, dict)
        )
        if has_cap_table_xrefs:
            try:
                from app.services.legal_cap_table_bridge import LegalCapTableBridge
                bridge = LegalCapTableBridge()
                bridge_result = bridge.build_from_documents(company_id, fund_id, document_id)
                if bridge_result.get("success"):
                    logger.info("[DOC_CAP_TABLE] Built document-derived cap table for company %s from document %s", company_id, document_id)
                    try:
                        from app.services.micro_skills.suggestion_emitter import emit_cap_table_suggestions
                        emit_cap_table_suggestions(
                            cap_table_result=bridge_result,
                            company_id=company_id,
                            fund_id=fund_id,
                            document_id=document_id,
                            document_name=document_name,
                        )
                    except Exception as e_emit:
                        logger.warning("[DOC_CAP_TABLE] Cap table suggestion emission failed: %s", e_emit)
            except Exception as e:
                logger.warning("[DOC_CAP_TABLE] Document-derived cap table failed for %s: %s", document_id, e, exc_info=True)

    # ── Path B: Synthetic cap table from funding signals (fallback) ──
    if extracted_data:
        has_funding_signal = (
            extracted_data.get("stage")
            or extracted_data.get("total_funding")
            or extracted_data.get("valuation_pre_money")
            or extracted_data.get("round")
        )
        if has_funding_signal:
            try:
                from app.services.pre_post_cap_table import PrePostCapTable
                from app.services.intelligent_gap_filler import IntelligentGapFiller

                gap_filler = IntelligentGapFiller()
                synthetic_rounds = gap_filler.generate_stage_based_funding_rounds(extracted_data)
                if not synthetic_rounds:
                    stage = extracted_data.get("stage") or extracted_data.get("round") or "Unknown"
                    amount = extracted_data.get("total_funding") or extracted_data.get("valuation_pre_money") or 0
                    if amount:
                        synthetic_rounds = [{"round": stage, "amount": amount}]

                if synthetic_rounds:
                    cap_data = {
                        "funding_rounds": synthetic_rounds,
                        "founders": [],
                        "is_yc": False,
                        "geography": extracted_data.get("geography", "Unknown"),
                    }
                    cap_service = PrePostCapTable()
                    cap_result = cap_service.calculate_full_cap_table_history(cap_data)

                    # Emit as suggestion instead of direct upsert — suggestions-first flow
                    try:
                        from app.services.micro_skills.suggestion_emitter import emit_cap_table_suggestions
                        emit_cap_table_suggestions(
                            cap_table_result=cap_result,
                            company_id=company_id,
                            fund_id=fund_id,
                            document_id=document_id,
                            document_name=document_name,
                        )
                        logger.info("[DOC_CAP_TABLE] Emitted cap table suggestion for company %s from document %s", company_id, document_id)
                    except Exception as e_emit:
                        logger.warning("[DOC_CAP_TABLE] Cap table suggestion emission failed: %s", e_emit)

                    # Founder ownership suggestion (already goes through suggestions)
                    from app.core.database import get_supabase_service
                    sb = get_supabase_service()
                    if sb:
                        client = sb.get_client() if hasattr(sb, 'get_client') else sb
                        if client:
                            founder_own = cap_result.get("founder_ownership")
                            if founder_own is not None:
                                client.table("pending_suggestions").upsert({
                                    "fund_id": fund_id or _UNLINKED,
                                    "company_id": company_id or _UNLINKED,
                                    "column_id": "founderOwnership",
                                    "suggested_value": {"value": founder_own},
                                    "source_service": "doc_cap_table",
                                    "reasoning": f"Founder ownership from document extraction: {founder_own:.1f}%",
                                }, on_conflict="fund_id,company_id,column_id").execute()
            except Exception as e:
                logger.warning("[DOC_CAP_TABLE] Cap table calculation failed for document %s: %s", document_id, e, exc_info=True)


def run_document_process(
    document_id: str,
    storage_path: str,
    document_type: str = "other",
    *,
    storage: DocumentBlobStorage,
    document_repo: DocumentMetadataRepo,
    company_id: Optional[str] = None,
    fund_id: Optional[str] = None,
    erp_category_hint: Optional[str] = None,
    erp_subcategory_hint: Optional[str] = None,
    force: bool = False,
) -> Dict[str, Any]:
    """
    Download document from storage, extract text, run structured extraction in-process,
    update metadata via repo. Progress is written to processing_summary at each step.
    Returns { "success": bool, "document_id": str, "result"?: dict, "error"?: str }.

    Pass force=True to re-extract even if the document is already completed or stuck
    in processing state.
    """
    doc_repo = document_repo
    tmp_path: Optional[str] = None

    # ── Idempotency guard: atomic claim prevents duplicate processing ──
    try:
        current_doc = doc_repo.get(document_id)
        if current_doc:
            current_status = current_doc.get("status")

            if force:
                logger.info(
                    "[DOC_PROCESS] Document %s force re-processing (was %s)",
                    document_id, current_status,
                )
                # Reset status so extraction runs fresh
                doc_repo.update(document_id, {"status": "processing"})
            elif current_status == "completed":
                extracted = current_doc.get("extracted_data") or {}
                doc_type_check = (current_doc.get("document_type") or document_type or "other").strip().lower()

                # Check if post-extraction pipeline was missed (e.g. transient DB failure).
                # If suggestions don't exist yet for this document, re-run the pipeline.
                needs_pipeline_rerun = False
                try:
                    from app.core.database import get_supabase_service
                    sb_check = get_supabase_service().get_client()
                    if sb_check and (company_id or fund_id):
                        existing = sb_check.table("pending_suggestions").select("id").eq(
                            "source_service", f"document:{document_id}"
                        ).limit(1).execute()
                        if not (existing.data if existing else []):
                            needs_pipeline_rerun = True
                            logger.warning(
                                "[DOC_PROCESS] Document %s is completed but has 0 pending_suggestions — re-running post-extraction pipeline",
                                document_id,
                            )
                except Exception as e:
                    logger.warning("[DOC_PROCESS] Could not check suggestions for %s: %s", document_id, e)

                if needs_pipeline_rerun and extracted:
                    try:
                        run_post_extraction_pipeline(
                            extracted_data=extracted,
                            document_id=document_id,
                            document_type=doc_type_check,
                            company_id=company_id or current_doc.get("company_id"),
                            fund_id=fund_id or current_doc.get("fund_id"),
                            document_name=Path(storage_path).stem if storage_path else "",
                            field_count=len([k for k, v in extracted.items() if v]),
                        )
                    except Exception as e:
                        logger.warning("[DOC_PROCESS] Post-extraction pipeline re-run failed for %s: %s", document_id, e, exc_info=True)

                logger.info(
                    "[DOC_PROCESS] Document %s already completed%s — returning cached result",
                    document_id, " (re-ran pipeline)" if needs_pipeline_rerun else "",
                )
                return {
                    "success": True,
                    "document_id": document_id,
                    "result": {
                        "extracted_data": extracted,
                        "processing_summary": current_doc.get("processing_summary", {}),
                    },
                }
            elif current_status == "processing":
                # Check for zombie: if processing_summary.updated_at is stale
                # (>5 min old), treat as stuck and re-process.
                summary = current_doc.get("processing_summary") or {}
                updated_at = summary.get("updated_at")
                is_zombie = False
                if updated_at:
                    try:
                        from datetime import datetime, timezone, timedelta
                        last_update = datetime.fromisoformat(updated_at.replace("Z", "+00:00"))
                        is_zombie = (datetime.now(timezone.utc) - last_update) > timedelta(minutes=5)
                    except Exception:
                        is_zombie = True  # Can't parse timestamp → assume stuck
                else:
                    is_zombie = True  # No timestamp at all → assume stuck

                if is_zombie:
                    logger.warning(
                        "[DOC_PROCESS] Document %s stuck in processing (zombie) — re-processing",
                        document_id,
                    )
                    doc_repo.update(document_id, {"status": "processing"})
                else:
                    logger.info(
                        "[DOC_PROCESS] Document %s already processing — skipping duplicate run",
                        document_id,
                    )
                    return {
                        "success": True,
                        "document_id": document_id,
                        "result": None,
                        "message": "Already processing",
                    }
            elif current_status == "pending":
                # Atomic claim: transition pending → processing in one DB call.
                claimed = doc_repo.claim_for_processing(document_id)
                if not claimed:
                    logger.info(
                        "[DOC_PROCESS] Document %s lost claim race — another caller is processing it",
                        document_id,
                    )
                    return {
                        "success": True,
                        "document_id": document_id,
                        "result": None,
                        "message": "Already claimed by another processor",
                    }
                logger.info("[DOC_PROCESS] Document %s claimed for processing", document_id)
    except Exception as e:
        logger.warning("[DOC_PROCESS] Idempotency check failed for %s: %s", document_id, e)
        # Continue processing — better to risk a duplicate than to block entirely

    def update_progress(step: str, message: str = "") -> None:
        try:
            doc_repo.update(
                document_id,
                {
                    "status": "processing",
                    "processing_summary": {
                        "step": step,
                        "message": message,
                        "updated_at": _iso_now(),
                    },
                },
            )
        except Exception as e:
            logger.warning("Could not update progress: %s", e)

    try:
        update_progress("downloading", "Downloading file from storage")
        content = storage.download(storage_path)
        if not content:
            doc_repo.update(
                document_id,
                {
                    "status": "failed",
                    "processing_summary": {"error": "Empty file from storage", "updated_at": _iso_now()},
                },
            )
            return {"success": False, "document_id": document_id, "error": "Empty file from storage"}

        suffix = Path(storage_path).suffix or ".pdf"
        fd, tmp_path = tempfile.mkstemp(suffix=suffix, prefix="doc_")
        try:
            os.write(fd, content)
        finally:
            os.close(fd)

        update_progress("extracting_text", "Extracting text from document")
        raw_text = _text_from_file(tmp_path, suffix)
        raw_text_preview = (raw_text[:5000] + "…") if len(raw_text) > 5000 else raw_text

        if not raw_text.strip():
            error_detail = "No text extracted from document"
            if suffix.lower().lstrip(".") == "pdf":
                error_detail = (
                    "No text extracted from PDF. "
                    "The file may be image-only (scanned) or password-protected. "
                    f"OCR available: {_OCR_AVAILABLE}."
                )
            doc_repo.update(
                document_id,
                {
                    "status": "failed",
                    "processing_summary": {
                        "error": error_detail,
                        "updated_at": _iso_now(),
                        "ocr_available": _OCR_AVAILABLE,
                    },
                },
            )
            logger.warning("[DOC_PROCESS] %s for document %s (path=%s)", error_detail, document_id, storage_path)
            return {"success": False, "document_id": document_id, "error": error_detail}

        update_progress("extracting_structured", "Running AI extraction")
        memo_context: Optional[str] = None
        if (document_type or "").strip().lower() != "investment_memo" and company_id:
            memo_context = _get_memo_context_for_company(doc_repo, company_id, fund_id)
        # Use asyncio.run() instead of manually creating/closing event loops.
        # asyncio.run() is cheaper, avoids resource leaks, and handles cleanup properly.
        extracted_data = asyncio.run(
            _extract_document_structured_async(
                raw_text, document_type or "other", memo_context=memo_context,
                erp_category_hint=erp_category_hint,
                erp_subcategory_hint=erp_subcategory_hint,
            )
        )

        # Check for extraction errors and log them
        extraction_error = extracted_data.get("_extraction_error")
        if extraction_error:
            logger.warning(
                "[DOC_PROCESS] Extraction completed with error for %s: %s",
                document_id, extraction_error,
            )

        # Count how many useful fields were extracted
        field_count = sum(
            1 for k, v in extracted_data.items()
            if v is not None and k not in ("_extraction_error", "value_explanations", "period_date")
            and not (isinstance(v, (list, dict)) and len(v) == 0)
            and not (isinstance(v, str) and not v.strip())
        )
        logger.info(
            "[DOC_PROCESS] Extracted %d non-empty fields from document %s",
            field_count, document_id,
        )

        processing_summary = {
            "step": "completed",
            "message": f"Extraction completed — {field_count} fields extracted"
                       + (f" (warning: {extraction_error})" if extraction_error else ""),
            "updated_at": _iso_now(),
            "fields_extracted": field_count,
            "text_length": len(raw_text),
        }
        update_payload: Dict[str, Any] = {
            "status": "completed",
            "processed_at": _iso_now(),
            "document_type": document_type or "other",
            "extracted_data": extracted_data,
            "issue_analysis": {},
            "comparables_analysis": {},
            "processing_summary": processing_summary,
            "raw_text_preview": raw_text_preview,
        }
        if company_id is not None:
            update_payload["company_id"] = company_id
        if fund_id is not None:
            update_payload["fund_id"] = fund_id
        doc_repo.update(document_id, update_payload)

        run_post_extraction_pipeline(
            extracted_data=extracted_data,
            document_id=document_id,
            document_type=document_type or "other",
            company_id=company_id,
            fund_id=fund_id,
            document_name=Path(storage_path).stem,
            field_count=field_count,
        )

        result = {
            "success": True,
            "extracted_data": extracted_data,
            "document_metadata": {"document_type": document_type or "other"},
            "processing_summary": processing_summary,
            "raw_text_preview": raw_text_preview,
            "issue_analysis": {},
            "comparables_analysis": {},
        }
        return {"success": True, "document_id": document_id, "result": result}
    except Exception as e:
        logger.exception("Document process failed: %s", e)
        try:
            doc_repo.update(
                document_id,
                {
                    "status": "failed",
                    "processing_summary": {"error": str(e), "updated_at": _iso_now()},
                },
            )
        except Exception:
            pass
        return {"success": False, "document_id": document_id, "error": str(e)}
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
