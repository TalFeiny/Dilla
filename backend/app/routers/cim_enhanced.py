"""
Enhanced CIM API endpoints with Google Docs-grade features
Supports real-time collaboration, charting, and rich document editing
"""

from fastapi import APIRouter, HTTPException, WebSocket, WebSocketDisconnect, UploadFile, File, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse
from typing import List, Dict, Any, Optional
from datetime import datetime, timedelta
import json
import asyncio
import uuid
from io import BytesIO
import base64

# Document generation libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.piecharts import Pie
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np

# Real-time collaboration
from ypy_websocket import WebsocketServer, YRoom
import y_py as Y

# Database and services
from app.core.database import supabase_service
# from app.services.supabase_service import SupabaseService
# from app.services.ai_service import AIService

router = APIRouter(prefix="/api/cim", tags=["Enhanced CIM"])

# Store active WebSocket connections for real-time collaboration
active_connections: Dict[str, List[WebSocket]] = {}
document_rooms: Dict[str, YRoom] = {}

# Initialize services
# supabase_service = SupabaseService()
# ai_service = AIService()

class CIMDocument:
    """Enhanced CIM document with versioning and collaboration"""
    
    def __init__(self, document_id: str):
        self.id = document_id
        self.ydoc = Y.YDoc()
        self.content = self.ydoc.get_text("content")
        self.metadata = self.ydoc.get_map("metadata")
        self.charts = self.ydoc.get_array("charts")
        self.comments = self.ydoc.get_array("comments")
        self.versions = []
        self.last_saved = datetime.now()
        self.collaborators = set()
    
    def add_version(self, author: str, changes: str):
        """Add a new version to the document history"""
        version = {
            "id": str(uuid.uuid4()),
            "timestamp": datetime.now().isoformat(),
            "author": author,
            "changes": changes,
            "content": self.content.to_json(),
            "stats": self.calculate_stats()
        }
        self.versions.append(version)
        return version
    
    def calculate_stats(self):
        """Calculate document statistics"""
        content_str = str(self.content)
        return {
            "characters": len(content_str),
            "words": len(content_str.split()),
            "charts": len(self.charts),
            "comments": len(self.comments)
        }

# Document storage (in production, use Redis or database)
documents: Dict[str, CIMDocument] = {}

@router.websocket("/ws/{document_id}")
async def websocket_endpoint(websocket: WebSocket, document_id: str):
    """WebSocket endpoint for real-time collaboration"""
    await websocket.accept()
    
    # Add connection to room
    if document_id not in active_connections:
        active_connections[document_id] = []
    active_connections[document_id].append(websocket)
    
    # Create or get document
    if document_id not in documents:
        documents[document_id] = CIMDocument(document_id)
    
    doc = documents[document_id]
    doc.collaborators.add(websocket)
    
    try:
        # Send initial document state
        await websocket.send_json({
            "type": "init",
            "content": doc.content.to_json() if hasattr(doc.content, 'to_json') else {},
            "metadata": dict(doc.metadata) if doc.metadata else {},
            "collaborators": len(doc.collaborators)
        })
        
        while True:
            # Receive updates from client
            data = await websocket.receive_json()
            
            if data["type"] == "update":
                # Apply update to document
                doc.content.apply_delta(data["delta"])
                
                # Broadcast to other clients
                for connection in active_connections[document_id]:
                    if connection != websocket:
                        await connection.send_json({
                            "type": "update",
                            "delta": data["delta"],
                            "user": data.get("user", "Anonymous")
                        })
            
            elif data["type"] == "cursor":
                # Broadcast cursor position
                for connection in active_connections[document_id]:
                    if connection != websocket:
                        await connection.send_json({
                            "type": "cursor",
                            "position": data["position"],
                            "user": data.get("user", "Anonymous")
                        })
            
            elif data["type"] == "comment":
                # Add comment
                comment = {
                    "id": str(uuid.uuid4()),
                    "user": data.get("user", "Anonymous"),
                    "content": data["content"],
                    "position": data.get("position"),
                    "timestamp": datetime.now().isoformat()
                }
                doc.comments.append(comment)
                
                # Broadcast comment
                for connection in active_connections[document_id]:
                    await connection.send_json({
                        "type": "comment",
                        "comment": comment
                    })
    
    except WebSocketDisconnect:
        active_connections[document_id].remove(websocket)
        doc.collaborators.discard(websocket)
        
        # Notify others about disconnection
        for connection in active_connections[document_id]:
            await connection.send_json({
                "type": "user_left",
                "collaborators": len(doc.collaborators)
            })

@router.post("/save")
async def save_document(request: Dict[str, Any]):
    """Save document with versioning"""
    document_id = request.get("document_id", str(uuid.uuid4()))
    content = request.get("content")
    author = request.get("author", "System")
    changes = request.get("changes", "Document saved")
    
    # Get or create document
    if document_id not in documents:
        documents[document_id] = CIMDocument(document_id)
    
    doc = documents[document_id]
    
    # Update content
    if content:
        doc.content = content
    
    # Add version
    version = doc.add_version(author, changes)
    doc.last_saved = datetime.now()
    
    # Save to database
    await supabase_service.save_cim_document(document_id, doc)
    
    return {
        "document_id": document_id,
        "version": version,
        "saved_at": doc.last_saved.isoformat()
    }

@router.get("/versions/{document_id}")
async def get_versions(document_id: str):
    """Get document version history"""
    if document_id not in documents:
        # Try to load from database
        doc_data = await supabase_service.get_cim_document(document_id)
        if not doc_data:
            raise HTTPException(status_code=404, detail="Document not found")
        documents[document_id] = CIMDocument.from_data(doc_data)
    
    doc = documents[document_id]
    return {"versions": doc.versions}

@router.post("/restore/{document_id}/{version_id}")
async def restore_version(document_id: str, version_id: str):
    """Restore document to a specific version"""
    if document_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc = documents[document_id]
    version = next((v for v in doc.versions if v["id"] == version_id), None)
    
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    
    # Restore content
    doc.content = version["content"]
    doc.add_version("System", f"Restored to version {version_id}")
    
    # Notify collaborators
    for connection in active_connections.get(document_id, []):
        await connection.send_json({
            "type": "version_restored",
            "version": version
        })
    
    return {"message": "Version restored successfully"}

@router.post("/export")
async def export_document(request: Dict[str, Any]):
    """Export CIM document to various formats"""
    content = request.get("content", "")
    format_type = request.get("format", "pdf")
    company_name = request.get("companyName", "Company")
    charts = request.get("charts", [])
    
    if format_type == "pdf":
        # Generate PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30
        )
        story.append(Paragraph(f"{company_name} - Confidential Information Memorandum", title_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Add content
        content_html = content.replace('\n', '<br/>')
        story.append(Paragraph(content_html, styles['Normal']))
        
        # Add charts if any
        for chart_config in charts:
            chart_image = generate_chart_image(chart_config)
            if chart_image:
                img = Image(chart_image, width=6*inch, height=4*inch)
                story.append(Spacer(1, 0.3*inch))
                story.append(img)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={company_name}_CIM.pdf"}
        )
    
    elif format_type == "docx":
        # Generate Word document
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(f"{company_name} - Confidential Information Memorandum", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph(content)
        
        # Add charts
        for chart_config in charts:
            chart_image = generate_chart_image(chart_config)
            if chart_image:
                doc.add_picture(chart_image, width=Inches(6))
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={company_name}_CIM.docx"}
        )
    
    elif format_type == "html":
        # Generate HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{company_name} - CIM</title>
            <style>
                body {{ font-family: Arial, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; }}
                h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                .chart {{ margin: 20px 0; text-align: center; }}
                .confidential {{ 
                    position: fixed; 
                    top: 50%; 
                    left: 50%; 
                    transform: translate(-50%, -50%) rotate(-45deg);
                    font-size: 120px;
                    color: rgba(255, 0, 0, 0.1);
                    z-index: -1;
                    font-weight: bold;
                }}
            </style>
        </head>
        <body>
            <div class="confidential">CONFIDENTIAL</div>
            <h1>{company_name} - Confidential Information Memorandum</h1>
            {content}
        </body>
        </html>
        """
        
        return StreamingResponse(
            BytesIO(html_content.encode()),
            media_type="text/html",
            headers={"Content-Disposition": f"attachment; filename={company_name}_CIM.html"}
        )
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported export format")

@router.post("/chart-suggestion")
async def suggest_chart_type(request: Dict[str, Any]):
    """AI-powered chart type suggestion based on data"""
    data = request.get("data", [])
    context = request.get("context", "")
    
    # Analyze data structure
    if not data:
        return {"type": "line", "message": "Default chart type"}
    
    # Get data characteristics
    sample = data[0] if data else {}
    keys = list(sample.keys())
    numeric_keys = [k for k in keys if isinstance(sample.get(k), (int, float))]
    
    # Use AI to suggest best chart type
    prompt = f"""
    Given this data structure:
    Keys: {keys}
    Numeric keys: {numeric_keys}
    Number of data points: {len(data)}
    Context: {context}
    
    Suggest the best chart type (line, bar, pie, area, scatter, radar, funnel).
    Also suggest which keys to use for x-axis and y-axis.
    """
    
    suggestion = await ai_service.get_chart_suggestion(prompt)
    
    return {
        "type": suggestion.get("type", "line"),
        "xAxis": suggestion.get("xAxis", keys[0] if keys else "x"),
        "yAxis": suggestion.get("yAxis", numeric_keys[:2] if numeric_keys else ["y"]),
        "reason": suggestion.get("reason", "Based on data structure analysis")
    }

@router.post("/generate-insights")
async def generate_insights(request: Dict[str, Any]):
    """Generate AI-powered insights from CIM data"""
    company_data = request.get("companyData", {})
    financial_data = request.get("financialData", {})
    market_data = request.get("marketData", {})
    
    insights = await ai_service.generate_cim_insights(
        company_data,
        financial_data,
        market_data
    )
    
    return {"insights": insights}

@router.get("/chart/render")
async def render_chart(config: str):
    """Render chart as image"""
    try:
        chart_config = json.loads(config)
        chart_image = generate_chart_image(chart_config)
        
        return StreamingResponse(
            chart_image,
            media_type="image/png"
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

def generate_chart_image(config: Dict[str, Any]) -> BytesIO:
    """Generate chart image from configuration"""
    chart_type = config.get("type", "line")
    data = pd.DataFrame(config.get("data", []))
    title = config.get("title", "Chart")
    x_axis = config.get("xAxis", data.columns[0] if len(data.columns) > 0 else "x")
    y_axis = config.get("yAxis", [data.columns[1]] if len(data.columns) > 1 else ["y"])
    
    # Create figure
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Generate chart based on type
    if chart_type == "line":
        for y in y_axis:
            if y in data.columns:
                ax.plot(data[x_axis], data[y], marker='o', label=y)
    
    elif chart_type == "bar":
        if len(y_axis) == 1 and y_axis[0] in data.columns:
            ax.bar(data[x_axis], data[y_axis[0]])
        else:
            # Multiple bars
            x = np.arange(len(data[x_axis]))
            width = 0.8 / len(y_axis)
            for i, y in enumerate(y_axis):
                if y in data.columns:
                    ax.bar(x + i * width, data[y], width, label=y)
    
    elif chart_type == "pie":
        if y_axis[0] in data.columns:
            ax.pie(data[y_axis[0]], labels=data[x_axis], autopct='%1.1f%%')
    
    elif chart_type == "scatter":
        if len(y_axis) >= 1 and y_axis[0] in data.columns:
            ax.scatter(data[x_axis], data[y_axis[0]])
    
    elif chart_type == "area":
        for y in y_axis:
            if y in data.columns:
                ax.fill_between(data[x_axis], data[y], alpha=0.5, label=y)
    
    # Styling
    ax.set_title(title, fontsize=16, fontweight='bold')
    ax.set_xlabel(x_axis, fontsize=12)
    if chart_type != "pie":
        ax.set_ylabel(", ".join(y_axis), fontsize=12)
        ax.legend()
    ax.grid(True, alpha=0.3)
    
    # Save to buffer
    buffer = BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
    buffer.seek(0)
    plt.close()
    
    return buffer

@router.post("/comments/add")
async def add_comment(request: Dict[str, Any]):
    """Add a comment to the document"""
    document_id = request.get("document_id")
    comment = request.get("comment")
    
    if document_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc = documents[document_id]
    comment["id"] = str(uuid.uuid4())
    comment["timestamp"] = datetime.now().isoformat()
    doc.comments.append(comment)
    
    # Broadcast to collaborators
    for connection in active_connections.get(document_id, []):
        await connection.send_json({
            "type": "new_comment",
            "comment": comment
        })
    
    return {"comment": comment}

@router.post("/comments/resolve/{comment_id}")
async def resolve_comment(comment_id: str, request: Dict[str, Any]):
    """Resolve a comment"""
    document_id = request.get("document_id")
    
    if document_id not in documents:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc = documents[document_id]
    
    # Find and update comment
    for comment in doc.comments:
        if comment.get("id") == comment_id:
            comment["resolved"] = True
            comment["resolved_at"] = datetime.now().isoformat()
            
            # Broadcast update
            for connection in active_connections.get(document_id, []):
                await connection.send_json({
                    "type": "comment_resolved",
                    "comment_id": comment_id
                })
            
            return {"message": "Comment resolved"}
    
    raise HTTPException(status_code=404, detail="Comment not found")

@router.get("/templates")
async def get_cim_templates():
    """Get available CIM templates"""
    templates = [
        {
            "id": "tech_saas",
            "name": "Technology SaaS",
            "description": "Template for SaaS companies",
            "sections": ["Executive Summary", "Company Overview", "Market Analysis", "Financial Performance", "Growth Strategy"]
        },
        {
            "id": "healthcare",
            "name": "Healthcare",
            "description": "Template for healthcare companies",
            "sections": ["Executive Summary", "Clinical Overview", "Market Opportunity", "Regulatory Status", "Financial Projections"]
        },
        {
            "id": "fintech",
            "name": "FinTech",
            "description": "Template for financial technology companies",
            "sections": ["Executive Summary", "Product Suite", "Regulatory Compliance", "Market Analysis", "Unit Economics"]
        }
    ]
    
    return {"templates": templates}

@router.post("/generate-from-template")
async def generate_from_template(request: Dict[str, Any]):
    """Generate CIM content from template"""
    template_id = request.get("template_id")
    company_data = request.get("company_data", {})
    
    # Generate content using AI
    content = await ai_service.generate_cim_from_template(template_id, company_data)
    
    return {"content": content}